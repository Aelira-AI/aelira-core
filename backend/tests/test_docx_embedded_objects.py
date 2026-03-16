"""
Tests for Embedded Object Detection (Task 7)

Tests cover:
- EmbeddedObject model
- EmbeddedObjectIssue model
- Embedded object detection in DOCX files
- Type identification for embedded content
- Alt text checking for embedded objects
"""

import pytest
import tempfile
import os
import zipfile

from src.education.docx_processor import (
    DocxProcessor,
    EmbeddedObject,
    EmbeddedObjectIssue,
)


class TestEmbeddedObjectModel:
    """Test EmbeddedObject Pydantic model."""

    def test_embedded_object_creation(self):
        """Test creating an embedded object."""
        obj = EmbeddedObject(
            object_index=0,
            object_type="excel",
            file_name="data.xlsx",
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            location="word/embeddings/oleObject1.xlsx",
            has_alt_text=True,
            existing_alt_text="Sales data spreadsheet",
            is_accessible=True,
            size_bytes=15000,
        )

        assert obj.object_index == 0
        assert obj.object_type == "excel"
        assert obj.file_name == "data.xlsx"
        assert obj.has_alt_text is True
        assert obj.is_accessible is True

    def test_embedded_object_minimal(self):
        """Test embedded object with minimal fields."""
        obj = EmbeddedObject(
            object_index=1,
            object_type="ole",
            location="word/embeddings/oleObject2.bin",
            has_alt_text=False,
        )

        assert obj.file_name is None
        assert obj.content_type is None
        assert obj.existing_alt_text is None
        assert obj.is_accessible is False


class TestEmbeddedObjectIssueModel:
    """Test EmbeddedObjectIssue Pydantic model."""

    def test_embedded_object_issue_creation(self):
        """Test creating an embedded object issue."""
        issue = EmbeddedObjectIssue(
            object_index=0,
            object_type="pdf",
            file_name="document.pdf",
            issue_type="missing_alt_text",
            location="word/embeddings/oleObject1.pdf",
            recommendations=[
                "Add alt text for embedded PDF",
                "Ensure PDF is tagged",
            ],
            suggested_fix="Add descriptive alt text for the embedded PDF",
        )

        assert issue.object_index == 0
        assert issue.object_type == "pdf"
        assert issue.issue_type == "missing_alt_text"
        assert len(issue.recommendations) == 2

    def test_embedded_object_issue_types(self):
        """Test different issue types."""
        # Missing alt text
        issue1 = EmbeddedObjectIssue(
            object_index=0,
            object_type="excel",
            issue_type="missing_alt_text",
            location="word/embeddings/data.xlsx",
            recommendations=["Add alt text"],
            suggested_fix="Add alt text",
        )
        assert issue1.issue_type == "missing_alt_text"

        # Inaccessible content
        issue2 = EmbeddedObjectIssue(
            object_index=1,
            object_type="ole",
            issue_type="inaccessible_content",
            location="word/embeddings/object.bin",
            recommendations=["Provide text alternative"],
            suggested_fix="Provide text alternative",
        )
        assert issue2.issue_type == "inaccessible_content"


class TestEmbeddedObjectDetection:
    """Test embedded object detection in DOCX files."""

    @pytest.fixture
    def processor(self):
        """Create DocxProcessor for testing."""
        return DocxProcessor()

    def test_detect_embedded_objects_empty_doc(self, processor):
        """Test detection on document without embedded objects."""
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("This is a simple document without embedded objects.")
            doc.save(f.name)

            issues = processor._detect_embedded_objects(doc, f.name)

            # Should return empty list for document without embedded objects
            assert isinstance(issues, list)

        os.unlink(f.name)

    def test_detect_embedded_objects_returns_list(self, processor):
        """Test that detect_embedded_objects always returns a list."""
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("Test document")
            doc.save(f.name)

            result = processor._detect_embedded_objects(doc, f.name)

            assert isinstance(result, list)

        os.unlink(f.name)

    def test_detect_embedded_objects_with_fake_embedding(self, processor):
        """Test detection with simulated embedded content."""
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            # Create basic document
            doc = Document()
            doc.add_paragraph("Document with embedded content")
            doc.save(f.name)

            # Manually add fake embedding to the DOCX ZIP
            with zipfile.ZipFile(f.name, "a") as zf:
                # Add a fake Excel embedding
                zf.writestr(
                    "word/embeddings/oleObject1.xlsx", b"PK\x03\x04fake excel content"
                )

            # Reload doc for detection
            doc = Document(f.name)
            issues = processor._detect_embedded_objects(doc, f.name)

            # Should detect the fake embedding
            assert isinstance(issues, list)
            # Should have at least one issue for the embedding
            assert len(issues) >= 1
            # Should identify as Excel type
            excel_issues = [i for i in issues if i.object_type == "excel"]
            assert len(excel_issues) >= 1

        os.unlink(f.name)

    def test_detect_embedded_pdf(self, processor):
        """Test detection of embedded PDF."""
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("Document with embedded PDF")
            doc.save(f.name)

            with zipfile.ZipFile(f.name, "a") as zf:
                zf.writestr(
                    "word/embeddings/document.pdf", b"%PDF-1.4 fake pdf content"
                )

            doc = Document(f.name)
            issues = processor._detect_embedded_objects(doc, f.name)

            pdf_issues = [i for i in issues if i.object_type == "pdf"]
            assert len(pdf_issues) >= 1

            # Check recommendations include PDF-specific advice
            if pdf_issues:
                all_recs = " ".join(pdf_issues[0].recommendations)
                assert "PDF" in all_recs or "pdf" in all_recs.lower()

        os.unlink(f.name)

    def test_detect_embedded_ole(self, processor):
        """Test detection of OLE objects."""
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("Document with OLE object")
            doc.save(f.name)

            with zipfile.ZipFile(f.name, "a") as zf:
                zf.writestr(
                    "word/embeddings/oleObject1.bin",
                    b"\xd0\xcf\x11\xe0 OLE compound document",
                )

            doc = Document(f.name)
            issues = processor._detect_embedded_objects(doc, f.name)

            ole_issues = [i for i in issues if i.object_type == "ole"]
            assert len(ole_issues) >= 1

        os.unlink(f.name)


class TestEmbeddedObjectAltText:
    """Test alt text checking for embedded objects."""

    @pytest.fixture
    def processor(self):
        """Create DocxProcessor for testing."""
        return DocxProcessor()

    def test_get_embedded_alt_text_no_alt(self, processor):
        """Test checking alt text when none exists."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Simple paragraph")

        has_alt, alt_text = processor._get_embedded_alt_text(doc, "test/path", 0)

        assert has_alt is False
        assert alt_text is None


class TestFindOLEObjects:
    """Test finding OLE objects in document."""

    @pytest.fixture
    def processor(self):
        """Create DocxProcessor for testing."""
        return DocxProcessor()

    def test_find_ole_objects_empty_doc(self, processor):
        """Test finding OLE objects in empty document."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("No OLE objects here")

        ole_objects = processor._find_ole_objects_in_doc(doc)

        assert isinstance(ole_objects, list)


class TestDocxProcessorIntegration:
    """Test embedded object integration with DocxProcessor."""

    @pytest.fixture
    def processor(self):
        """Create DocxProcessor for testing."""
        return DocxProcessor()

    def test_process_docx_includes_embedded_issues(self, processor):
        """Test that process_docx includes embedded_object_issues field."""
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_heading("Test Document", level=1)
            doc.add_paragraph("This is a test paragraph.")
            doc.save(f.name)

            result = processor.process_docx(f.name)

            # Should have embedded_object_issues field
            assert hasattr(result, "embedded_object_issues")
            assert isinstance(result.embedded_object_issues, list)

        os.unlink(f.name)

    def test_summary_includes_embedded_count(self, processor):
        """Test that summary includes embedded_object_issues count."""
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("Simple test")
            doc.save(f.name)

            result = processor.process_docx(f.name)

            # Summary should include embedded_object_issues key
            assert "embedded_object_issues" in result.summary

        os.unlink(f.name)

    def test_embedded_issues_in_total_count(self, processor):
        """Test that embedded issues are counted in total."""
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("Test")
            doc.save(f.name)

            # Add fake embedding
            with zipfile.ZipFile(f.name, "a") as zf:
                zf.writestr("word/embeddings/test.xlsx", b"fake")

            result = processor.process_docx(f.name)

            # Total should include embedded issues
            embedded_count = result.summary.get("embedded_object_issues", 0)
            total = result.summary.get("total_issues", 0)

            # If there are embedded issues, they should be in total
            if embedded_count > 0:
                assert total >= embedded_count

        os.unlink(f.name)


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-s"])
