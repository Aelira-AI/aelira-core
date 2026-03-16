#!/usr/bin/env python3
"""
Test script for the Auto-Remediation Engine.

This script tests the remediation framework by:
1. Creating test documents with accessibility issues
2. Running the appropriate remediator
3. Verifying fixes were applied
"""

import os
import sys
import tempfile
from pathlib import Path

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from openpyxl import Workbook

def create_test_word_document():
    """Create a Word document with accessibility issues."""
    doc = Document()

    # Add content without proper heading structure
    p = doc.add_paragraph("Introduction to the Course")  # Should be Heading
    p.runs[0].bold = True
    p.runs[0].font.size = 180000  # 18pt - heading-like

    doc.add_paragraph("This is the introduction to our course materials.")

    # Fake bullets (accessibility issue)
    doc.add_paragraph("• First fake bullet item")
    doc.add_paragraph("• Second fake bullet item")
    doc.add_paragraph("• Third fake bullet item")

    # Another heading-like paragraph
    p2 = doc.add_paragraph("Course Objectives")
    p2.runs[0].bold = True

    # Non-descriptive link text
    doc.add_paragraph("For more information, click here: https://example.com")

    # Add a table without headers
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Score"
    table.cell(0, 2).text = "Grade"
    table.cell(1, 0).text = "John"
    table.cell(1, 1).text = "85"
    table.cell(1, 2).text = "B"
    table.cell(2, 0).text = "Jane"
    table.cell(2, 1).text = "92"
    table.cell(2, 2).text = "A"

    # Save to temp file
    temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(temp_file.name)
    return temp_file.name

def create_test_excel_spreadsheet():
    """Create an Excel spreadsheet with accessibility issues."""
    wb = Workbook()

    # Generic sheet name (issue)
    ws = wb.active
    ws.title = "Sheet1"  # Generic name - should be renamed

    # Data without proper headers
    ws['A1'] = 'Product'
    ws['B1'] = 'Price'
    ws['C1'] = 'Quantity'
    ws['A2'] = 'Widget'
    ws['B2'] = 25.99
    ws['C2'] = 100
    ws['A3'] = 'Gadget'
    ws['B3'] = 49.99
    ws['C3'] = 50

    # Save to temp file
    temp_file = tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False)
    wb.save(temp_file.name)
    return temp_file.name

def test_docx_remediator():
    """Test the Word document remediator."""
    print("\n" + "="*60)
    print("Testing DocxRemediator")
    print("="*60)

    from src.education.remediation import DocxRemediator, RemediationConfig

    # Create test document
    file_path = create_test_word_document()
    print(f"Created test document: {file_path}")

    # Create sample issues (simulating scan results)
    issues = [
        {
            'type': 'list',
            'severity': 'medium',
            'description': 'Fake bullet list detected',
            'paragraph_indices': [2, 3, 4],
            'is_fake_list': True
        },
        {
            'type': 'heading',
            'severity': 'high',
            'description': 'Missing heading structure',
            'paragraph_index': 0,
            'suggested_level': 1
        },
        {
            'type': 'table',
            'severity': 'medium',
            'description': 'Table missing headers',
            'table_index': 0,
            'has_data_rows': True
        },
        {
            'type': 'language',
            'severity': 'low',
            'description': 'Document language not specified'
        }
    ]

    # Create config
    config = RemediationConfig(
        use_ai=False,  # No AI for testing
        verify_fixes=True,
        create_backup=True
    )

    # Create remediator and run
    try:
        remediator = DocxRemediator(
            file_path=file_path,
            issues=issues,
            config=config
        )

        result = remediator.remediate()

        print(f"\nRemediation Result:")
        print(f"  Success: {result.success}")
        print(f"  Total Issues: {result.total_issues}")
        print(f"  Fixed: {result.fixed_count}")
        print(f"  Manual: {result.manual_count}")
        print(f"  Failed: {result.failed_count}")
        print(f"  Duration: {result.duration_seconds:.2f}s")
        print(f"  Output: {result.output_file}")

        if result.fixed_issues:
            print(f"\n  Fixed Issues:")
            for fixed in result.fixed_issues:
                print(f"    - {fixed.category.value}: {fixed.description[:50]}...")

        if result.manual_issues:
            print(f"\n  Manual Issues (need human review):")
            for manual in result.manual_issues:
                print(f"    - {manual.category.value}: {manual.reason}")

        if result.warnings:
            print(f"\n  Warnings:")
            for warning in result.warnings:
                print(f"    - {warning}")

        # Cleanup
        os.unlink(file_path)
        if result.output_file and os.path.exists(result.output_file):
            print(f"\n  Remediated file saved at: {result.output_file}")

        return result.success

    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()
        os.unlink(file_path)
        return False

def test_xlsx_remediator():
    """Test the Excel spreadsheet remediator."""
    print("\n" + "="*60)
    print("Testing XlsxRemediator")
    print("="*60)

    from src.education.remediation import XlsxRemediator, RemediationConfig

    # Create test spreadsheet
    file_path = create_test_excel_spreadsheet()
    print(f"Created test spreadsheet: {file_path}")

    # Create sample issues
    issues = [
        {
            'type': 'sheet',
            'severity': 'medium',
            'description': 'Generic sheet name "Sheet1"',
            'sheet_name': 'Sheet1',
            'sheet_index': 0
        },
        {
            'type': 'table',
            'severity': 'high',
            'description': 'Table headers not defined',
            'sheet_name': 'Sheet1',
            'has_data': True
        },
        {
            'type': 'navigation',
            'severity': 'low',
            'description': 'Headers not frozen for navigation',
            'sheet_name': 'Sheet1'
        }
    ]

    config = RemediationConfig(
        use_ai=False,
        verify_fixes=True,
        create_backup=True
    )

    try:
        remediator = XlsxRemediator(
            file_path=file_path,
            issues=issues,
            config=config
        )

        result = remediator.remediate()

        print(f"\nRemediation Result:")
        print(f"  Success: {result.success}")
        print(f"  Total Issues: {result.total_issues}")
        print(f"  Fixed: {result.fixed_count}")
        print(f"  Manual: {result.manual_count}")
        print(f"  Failed: {result.failed_count}")
        print(f"  Duration: {result.duration_seconds:.2f}s")
        print(f"  Output: {result.output_file}")

        if result.fixed_issues:
            print(f"\n  Fixed Issues:")
            for fixed in result.fixed_issues:
                print(f"    - {fixed.category.value}: {fixed.description[:50]}...")

        # Cleanup
        os.unlink(file_path)
        if result.output_file and os.path.exists(result.output_file):
            print(f"\n  Remediated file saved at: {result.output_file}")

        return result.success

    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()
        os.unlink(file_path)
        return False

def test_remediator_factory():
    """Test the get_remediator_for_file function."""
    print("\n" + "="*60)
    print("Testing get_remediator_for_file")
    print("="*60)

    from src.education.remediation import get_remediator_for_file

    test_cases = [
        ('document.docx', 'DocxRemediator'),
        ('presentation.pptx', 'PptxRemediator'),
        ('report.pdf', 'PdfRemediator'),
        ('data.xlsx', 'XlsxRemediator'),
        ('unknown.txt', None),
    ]

    all_passed = True
    for filename, expected in test_cases:
        remediator_class = get_remediator_for_file(filename)

        if expected is None:
            passed = remediator_class is None
        else:
            passed = remediator_class is not None and remediator_class.__name__ == expected

        status = "✓" if passed else "✗"
        actual = remediator_class.__name__ if remediator_class else None
        print(f"  {status} {filename} -> {actual} (expected: {expected})")

        if not passed:
            all_passed = False

    return all_passed

def main():
    """Run all remediation tests."""
    print("\n" + "="*60)
    print("Aelira Auto-Remediation Engine - Test Suite")
    print("="*60)

    results = []

    # Test factory function
    results.append(("Factory Function", test_remediator_factory()))

    # Test Word remediator
    results.append(("DocxRemediator", test_docx_remediator()))

    # Test Excel remediator
    results.append(("XlsxRemediator", test_xlsx_remediator()))

    # Summary
    print("\n" + "="*60)
    print("Test Summary")
    print("="*60)

    passed = sum(1 for _, r in results if r)
    total = len(results)

    for name, result in results:
        status = "✓ PASSED" if result else "✗ FAILED"
        print(f"  {status}: {name}")

    print(f"\nTotal: {passed}/{total} tests passed")

    return passed == total

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
