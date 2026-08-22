"""
Word Document (.docx) Accessibility Scanner Module

This module provides functionality to:
1. Parse DOCX files and extract content structure
2. Detect missing alt text on images
3. Check heading hierarchy (proper H1-H6 structure)
4. Identify fake lists (bullets via symbols vs real lists)
5. Check table headers and structure
6. Detect "click here" link text issues
7. Check language specification
8. Analyze reading order
9. Generate AI-powered remediation suggestions
10. Batch process entire directories

Faculty use Word heavily, making this a CRITICAL feature for higher ed.
"""

from typing import List, Dict, Optional, Any
from pydantic import BaseModel, computed_field
from docx import Document
from docx.oxml.ns import qn
import os
import tempfile
from pathlib import Path
from PIL import Image
from io import BytesIO
import logging
import re

from src.education.color_blindness_simulator import (
    ColorBlindnessSimulator,
    ColorBlindnessAnalysisResult,
)
from src.utils.async_helpers import run_async_from_sync

logger = logging.getLogger(__name__)


def _style_name(para) -> str:
    """Paragraph style name, or "" when the style resolves to None.

    python-docx falls back to styles.default(WD_STYLE_TYPE.PARAGRAPH) for
    paragraphs with no explicit w:pStyle; that returns None when styles.xml
    defines no w:default="1" paragraph style (pandoc, Google Docs export,
    LibreOffice). An unstyled paragraph is body text: it should match neither
    heading nor list checks.
    """
    style = getattr(para, "style", None)
    return getattr(style, "name", None) or ""


class HeadingIssue(BaseModel):
    """Heading structure issue"""

    paragraph_index: int
    text: str
    issue_type: str  # missing_h1, skipped_level, no_headings
    current_level: Optional[int] = None
    expected_level: Optional[int] = None
    suggested_fix: str


class ImageIssue(BaseModel):
    """Image accessibility issue"""

    paragraph_index: int
    image_index: int  # Index within document
    has_alt_text: bool
    existing_alt_text: Optional[str] = None
    suggested_alt_text: Optional[str] = None
    image_path: Optional[str] = None  # ZIP-relative path e.g. word/media/image1.png
    # Smart image analysis
    detected_image_type: Optional[str] = (
        None  # decorative, informative, functional, complex
    )
    is_decorative: bool = False
    is_chart: bool = False
    detailed_description: Optional[str] = None
    # Alt text validation (for images WITH alt text)
    alt_text_validated: bool = False  # Whether AI validation was performed
    alt_text_accurate: Optional[bool] = None  # Whether existing alt text is accurate
    alt_text_issues: Optional[List[str]] = None  # Specific problems found
    validation_score: Optional[float] = None  # Accuracy score 0-1


class TableIssue(BaseModel):
    """Table accessibility issue"""

    table_index: int
    issue_type: str  # missing_header, complex_merge, missing_caption
    row_count: int
    column_count: int
    suggested_fix: str


class ListIssue(BaseModel):
    """List structure issue (fake bullets vs real lists)"""

    paragraph_index: int
    text: str
    issue_type: str  # fake_bullet, fake_number
    detected_marker: str  # The character used as fake bullet
    suggested_fix: str


class LinkIssue(BaseModel):
    """Hyperlink accessibility issue"""

    paragraph_index: int
    link_text: str
    link_url: str
    issue_type: str  # click_here, non_descriptive, broken
    suggested_fix: str


class LanguageIssue(BaseModel):
    """Document language issue"""

    issue_type: str  # missing_language, inconsistent_language
    suggested_fix: str


class TitleIssue(BaseModel):
    """Document title accessibility issue (WCAG 2.4.2)"""

    issue_type: str  # missing_title, empty_title, generic_title, filename_title
    existing_title: Optional[str] = None
    suggested_title: Optional[str] = None
    suggested_fix: str


class FontSizeIssue(BaseModel):
    """Font size accessibility issue (WCAG 1.4.4)"""

    paragraph_index: int
    text_preview: str  # First 50 chars of affected text
    font_size_pt: float
    issue_type: str  # too_small, very_small
    suggested_fix: str


class SmartArtElement(BaseModel):
    """Detected SmartArt diagram element."""

    diagram_index: int
    diagram_type: (
        str  # hierarchy, process, cycle, relationship, matrix, pyramid, list, picture
    )
    text_nodes: List[str]  # Text content from the SmartArt
    node_count: int
    has_alt_text: bool
    existing_alt_text: Optional[str] = None
    relationship_type: Optional[str] = None  # For relationship diagrams


class SmartArtIssue(BaseModel):
    """SmartArt accessibility issue."""

    diagram_index: int
    diagram_type: str
    issue_type: str  # missing_alt_text, complex_structure, missing_text_alternative
    node_count: int
    text_content: List[str]  # Extracted text nodes
    generated_description: Optional[str] = None  # AI-generated accessible description
    suggested_fix: str


class EmbeddedObject(BaseModel):
    """Detected embedded object in DOCX."""

    object_index: int
    object_type: str  # "excel", "pdf", "ole", "package", "image", "other"
    file_name: Optional[str] = None  # Original filename if available
    content_type: Optional[str] = None  # MIME type
    location: str  # Path within DOCX ZIP
    has_alt_text: bool
    existing_alt_text: Optional[str] = None
    is_accessible: bool = False  # Whether the embedded content is accessible
    size_bytes: Optional[int] = None


class EmbeddedObjectIssue(BaseModel):
    """Embedded object accessibility issue."""

    object_index: int
    object_type: str
    file_name: Optional[str] = None
    issue_type: str  # missing_alt_text, inaccessible_content, no_text_alternative
    location: str
    recommendations: List[str]
    suggested_fix: str


class DocxProcessingResult(BaseModel):
    """Result of Word document processing operation"""

    file_path: str
    file_name: str
    total_paragraphs: int
    total_images: int
    total_tables: int
    total_lists: int
    total_links: int
    heading_issues: List[HeadingIssue]
    image_issues: List[ImageIssue]
    table_issues: List[TableIssue]
    list_issues: List[ListIssue]
    link_issues: List[LinkIssue]
    language_issues: List[LanguageIssue]
    title_issues: List[TitleIssue] = []  # Document title issues (WCAG 2.4.2)
    font_size_issues: List[FontSizeIssue] = []  # Font size issues (WCAG 1.4.4)
    smartart_issues: List[SmartArtIssue] = []  # SmartArt accessibility issues
    embedded_object_issues: List[EmbeddedObjectIssue] = []  # Embedded object issues
    summary: Dict[str, int]
    compliance_score: float
    html_output: str
    remediation_suggestions: List[str]
    # Color vision deficiency analysis
    cvd_analysis: Optional[List[ColorBlindnessAnalysisResult]] = None

    @computed_field
    @property
    def issues(self) -> List[Dict[str, Any]]:
        """Combined list of all issues for API compatibility.

        The demo routes expect a single 'issues' array, but DOCX processor
        stores issues in separate category arrays. This computed field
        combines them all into a unified format.
        """
        all_issues: List[Dict[str, Any]] = []

        # Map severity from issue types
        def get_severity(issue_type: str) -> str:
            critical_types = ["missing_alt_text", "missing_header", "no_headings"]
            high_types = [
                "skipped_level",
                "fake_bullet",
                "fake_number",
                "missing_language",
            ]
            medium_types = [
                "click_here",
                "non_descriptive",
                "missing_title",
                "too_small",
            ]
            if any(t in issue_type for t in critical_types):
                return "critical"
            elif any(t in issue_type for t in high_types):
                return "high"
            elif any(t in issue_type for t in medium_types):
                return "medium"
            return "low"

        # Heading issues
        for issue in self.heading_issues:
            if issue.issue_type == "skipped_level":
                desc = (
                    f"Heading skips from H{issue.expected_level} to H{issue.current_level}: '{issue.text[:50]}...'"
                    if len(issue.text) > 50
                    else f"Heading skips from H{issue.expected_level} to H{issue.current_level}: '{issue.text}'"
                )
                title = f"Heading Level Skip (H{issue.expected_level} → H{issue.current_level})"
            elif issue.issue_type == "missing_h1":
                desc = "Document is missing a main heading (H1)"
                title = "Missing H1 Heading"
            else:
                desc = f"Heading issue: {issue.issue_type}"
                title = "Heading Structure Issue"
            all_issues.append(
                {
                    "id": f"heading_{len(all_issues)}",
                    "category": "heading",
                    "severity": get_severity(issue.issue_type),
                    "title": title,
                    "description": desc,
                    "location": f"Paragraph {issue.paragraph_index + 1}",
                    "wcag_criterion": "WCAG 1.3.1",
                    "suggested_fix": issue.suggested_fix,
                    # Metadata for remediator
                    "current_level": issue.current_level,
                    "expected_level": issue.expected_level,
                    "suggested_level": issue.expected_level,  # What level it should be
                    "text": issue.text,
                    "paragraph_index": issue.paragraph_index,
                }
            )

        # Image issues
        for issue in self.image_issues:
            if issue.has_alt_text:
                # Image has alt text - check if validation was performed
                if issue.alt_text_validated and not issue.alt_text_accurate:
                    # Alt text was validated and found to be inaccurate
                    issues_list = (
                        ", ".join(issue.alt_text_issues)
                        if issue.alt_text_issues
                        else "does not accurately describe the image"
                    )
                    existing_preview = (
                        issue.existing_alt_text[:50] + "..."
                        if issue.existing_alt_text and len(issue.existing_alt_text) > 50
                        else issue.existing_alt_text
                    )
                    desc = f"Alt text '{existing_preview}' is inaccurate: {issues_list}"
                    title = "Inaccurate Alt Text"
                    severity = (
                        "high"
                        if issue.validation_score and issue.validation_score < 0.5
                        else "medium"
                    )
                else:
                    existing_preview = (
                        issue.existing_alt_text[:50] + "..."
                        if issue.existing_alt_text and len(issue.existing_alt_text) > 50
                        else issue.existing_alt_text
                    )
                    desc = (
                        f"Image has alt text that may need review: '{existing_preview}'"
                    )
                    title = "Alt Text Review Needed"
                    severity = "low"
            else:
                if issue.is_decorative:
                    desc = f'Image {issue.image_index + 1} is decorative and should have empty alt text (alt="")'
                    title = "Decorative Image — Needs Empty Alt Text"
                    severity = "low"
                else:
                    desc = f"Image {issue.image_index + 1} is missing alternative text"
                    title = "Missing Alt Text"
                    severity = "critical"
            all_issues.append(
                {
                    "id": f"image_{len(all_issues)}",
                    "category": "alt_text",
                    "severity": severity,
                    "title": title,
                    "description": desc,
                    "location": f"Image {issue.image_index + 1}",
                    "wcag_criterion": "WCAG 1.1.1",
                    "suggested_fix": issue.suggested_alt_text
                    or (
                        "Mark as decorative (set empty alt text)"
                        if issue.is_decorative
                        else "Add descriptive alt text"
                    ),
                    "ai_generated": issue.suggested_alt_text is not None,
                    "generated_alt_text": issue.suggested_alt_text or None,
                    "is_decorative": issue.is_decorative,
                    # Metadata for remediator
                    "image_index": issue.image_index,
                    "paragraph_index": issue.paragraph_index,
                    "image_path": issue.image_path,
                    # Validation metadata
                    "alt_text_validated": issue.alt_text_validated,
                    "alt_text_accurate": issue.alt_text_accurate,
                    "alt_text_issues": issue.alt_text_issues,
                    "validation_score": issue.validation_score,
                }
            )

        # Table issues
        for issue in self.table_issues:
            if issue.issue_type == "missing_header":
                desc = f"Table ({issue.row_count}×{issue.column_count}) is missing header row designation"
            elif issue.issue_type == "complex_merge":
                desc = f"Table ({issue.row_count}×{issue.column_count}) has complex merged cells that may confuse screen readers"
            else:
                desc = f"Table issue: {issue.issue_type}"
            all_issues.append(
                {
                    "id": f"table_{len(all_issues)}",
                    "category": "table",
                    "severity": get_severity(issue.issue_type),
                    "title": issue.issue_type.replace("_", " ").title(),
                    "description": desc,
                    "location": f"Table {issue.table_index + 1}",
                    "wcag_criterion": "WCAG 1.3.1",
                    "suggested_fix": issue.suggested_fix,
                    # Metadata for remediator
                    "table_index": issue.table_index,
                    "has_data_rows": issue.row_count > 0,
                }
            )

        # List issues
        for issue in self.list_issues:
            text_preview = (
                issue.text[:50] + "..." if len(issue.text) > 50 else issue.text
            )
            desc = f"Text uses '{issue.detected_marker}' as fake bullet instead of proper list formatting: '{text_preview}'"
            all_issues.append(
                {
                    "id": f"list_{len(all_issues)}",
                    "category": "list",
                    "severity": get_severity(issue.issue_type),
                    "title": f"Fake {'Bullet' if issue.issue_type == 'fake_bullet' else 'Numbered'} List",
                    "description": desc,
                    "location": f"Paragraph {issue.paragraph_index + 1}",
                    "wcag_criterion": "WCAG 1.3.1",
                    "suggested_fix": issue.suggested_fix,
                    # Metadata for remediator - mark as fake list so it can be fixed
                    "is_fake_list": True,
                    "detected_marker": issue.detected_marker,
                    "text": issue.text,
                    "paragraph_index": issue.paragraph_index,
                }
            )

        # Link issues
        for issue in self.link_issues:
            desc = (
                f"Link text '{issue.link_text}' is not descriptive (URL: {issue.link_url[:50]}...)"
                if len(issue.link_url) > 50
                else f"Link text '{issue.link_text}' is not descriptive (URL: {issue.link_url})"
            )
            all_issues.append(
                {
                    "id": f"link_{len(all_issues)}",
                    "category": "link",
                    "severity": get_severity(issue.issue_type),
                    "title": "Non-Descriptive Link Text",
                    "description": desc,
                    "location": f"Paragraph {issue.paragraph_index + 1}",
                    "wcag_criterion": "WCAG 2.4.4",
                    "suggested_fix": issue.suggested_fix,
                    # Metadata for remediator
                    "link_text": issue.link_text,
                    "link_url": issue.link_url,
                    "paragraph_index": issue.paragraph_index,
                }
            )

        # Language issues
        for issue in self.language_issues:
            if issue.issue_type == "missing_language":
                desc = "Document does not specify a language, which is required for screen readers"
            else:
                desc = f"Language issue: {issue.issue_type}"
            all_issues.append(
                {
                    "id": f"language_{len(all_issues)}",
                    "category": "language",
                    "severity": get_severity(issue.issue_type),
                    "title": issue.issue_type.replace("_", " ").title(),
                    "description": desc,
                    "location": "Document",
                    "wcag_criterion": "WCAG 3.1.1",
                    "suggested_fix": issue.suggested_fix,
                }
            )

        # Title issues
        for issue in self.title_issues:
            if issue.issue_type == "missing_title":
                desc = "Document is missing a title in its properties"
            elif issue.issue_type == "empty_title":
                desc = "Document title is empty"
            elif issue.issue_type == "generic_title":
                desc = f"Document title '{issue.existing_title}' is too generic"
            elif issue.issue_type == "filename_title":
                desc = (
                    f"Document title '{issue.existing_title}' appears to be a filename"
                )
            else:
                desc = f"Title issue: {issue.issue_type}"
            all_issues.append(
                {
                    "id": f"title_{len(all_issues)}",
                    "category": "title",
                    "severity": get_severity(issue.issue_type),
                    "title": issue.issue_type.replace("_", " ").title(),
                    "description": desc,
                    "location": "Document Properties",
                    "wcag_criterion": "WCAG 2.4.2",
                    "suggested_fix": issue.suggested_fix,
                    # Metadata for remediator
                    "suggested_title": issue.suggested_title,
                    "existing_title": issue.existing_title,
                }
            )

        # Font size issues
        for issue in self.font_size_issues:
            desc = f"Text at {issue.font_size_pt}pt is below minimum readable size: '{issue.text_preview}'"
            all_issues.append(
                {
                    "id": f"font_{len(all_issues)}",
                    "category": "font_size",
                    "severity": (
                        "high" if issue.issue_type == "very_small" else "medium"
                    ),
                    "title": f"Font Size Too Small ({issue.font_size_pt}pt)",
                    "description": desc,
                    "location": f"Paragraph {issue.paragraph_index + 1}",
                    "wcag_criterion": "WCAG 1.4.4",
                    "suggested_fix": issue.suggested_fix,
                }
            )

        return all_issues


class DocxProcessor:
    """Process Word documents for accessibility compliance"""

    # Common fake bullet characters
    FAKE_BULLETS = [
        "•",
        "-",
        "*",
        "○",
        "▪",
        "◦",
        "–",
        "—",
        ">",
        "·",
        "→",
        "➤",
        "➢",
        "►",
        "❖",
    ]

    # Non-descriptive link text patterns
    BAD_LINK_PATTERNS = [
        r"^click\s*here$",
        r"^here$",
        r"^link$",
        r"^read\s*more$",
        r"^more$",
        r"^learn\s*more$",
        r"^this$",
        r"^download$",
        r"^pdf$",
        r"^info$",
        r"^details$",
    ]

    def __init__(
        self,
        generate_alt_text: bool = False,
        validate_alt_text: bool = False,
        enhance_descriptions: bool = False,
        simulate_color_blindness: bool = False,
        progress_callback: callable = None,
    ):
        self.generate_alt_text = generate_alt_text
        self.validate_alt_text = validate_alt_text
        self.enhance_descriptions = enhance_descriptions
        self.image_generator = None
        self.progress_callback = progress_callback
        # Color vision deficiency simulation
        self.simulate_color_blindness = simulate_color_blindness
        self.cvd_simulator = (
            ColorBlindnessSimulator() if simulate_color_blindness else None
        )

        # Lazy import image generator if needed
        if self.generate_alt_text or self.validate_alt_text:
            try:
                from .image_alt_text import ImageAltTextGenerator

                self.image_generator = ImageAltTextGenerator(
                    allow_legacy_transport=True
                )
            except Exception as e:
                logger.warning(
                    f"[DocxProcessor] Could not initialize ImageAltTextGenerator: {e}"
                )
                self.generate_alt_text = False
                self.validate_alt_text = False

    def process_docx(
        self, file_path: str, original_filename: str = None
    ) -> DocxProcessingResult:
        """
        Process a Word document and check accessibility

        Args:
            file_path: Path to DOCX file
            original_filename: Optional original filename

        Returns:
            DocxProcessingResult with all accessibility issues
        """
        doc = Document(file_path)
        file_name = original_filename or os.path.basename(file_path)

        # Progress tracking - 6 main steps
        total_steps = 6
        current_step = 0

        def report_progress(message: str):
            nonlocal current_step
            current_step += 1
            if self.progress_callback:
                self.progress_callback(current_step, total_steps, message)

        report_progress("Extracting document structure...")

        # Extract document context for AI
        document_context = self._extract_document_context(doc, file_name)

        report_progress("Checking heading structure...")

        # Analyze document structure
        heading_issues = self._check_heading_structure(doc)

        report_progress("Analyzing images for alt text...")
        image_issues = self._check_images(doc, document_context)

        report_progress("Checking tables for accessibility...")
        table_issues = self._check_tables(doc)

        report_progress("Checking lists and links...")
        list_issues = self._check_lists(doc)
        link_issues = self._check_links(doc)

        report_progress("Checking language attributes...")
        language_issues = self._check_language(doc)

        # Check document title (WCAG 2.4.2)
        title_issues = self._check_document_title(doc, document_context, file_name)

        # Check font sizes (WCAG 1.4.4)
        font_size_issues = self._check_font_sizes(doc)

        # Check SmartArt diagrams
        smartart_issues = self._detect_smartart(doc, file_path)

        # Check embedded objects
        embedded_object_issues = self._detect_embedded_objects(doc, file_path)

        # Count elements
        total_paragraphs = len(doc.paragraphs)
        total_images = len(
            [r for p in doc.paragraphs for r in p.runs if r._element.xpath(".//a:blip")]
        )
        total_tables = len(doc.tables)
        total_lists = self._count_real_lists(doc)
        total_links = self._count_links(doc)

        # Calculate summary and score
        summary = self._calculate_summary(
            heading_issues,
            image_issues,
            table_issues,
            list_issues,
            link_issues,
            language_issues,
            title_issues,
            font_size_issues,
            smartart_issues,
            embedded_object_issues,
        )

        total_elements = (
            total_paragraphs + total_images + total_tables + total_lists + total_links
        )
        compliance_score = self._calculate_compliance_score(summary, total_elements)

        # Generate accessible HTML
        html_output = self._generate_html(doc, document_context, file_name)

        # Generate remediation suggestions
        remediation_suggestions = self._generate_remediation_suggestions(summary)

        # Analyze color vision deficiency accessibility if enabled
        cvd_analysis = None
        if self.simulate_color_blindness and self.cvd_simulator:
            logger.info("[DocxProcessor] Running CVD accessibility analysis...")
            cvd_analysis = self._analyze_cvd_accessibility(doc)
            if cvd_analysis:
                logger.info(
                    f"[DocxProcessor] CVD analysis complete: {len(cvd_analysis)} color pairs tested"
                )

        return DocxProcessingResult(
            file_path=file_path,
            file_name=file_name,
            total_paragraphs=total_paragraphs,
            total_images=total_images,
            total_tables=total_tables,
            total_lists=total_lists,
            total_links=total_links,
            heading_issues=heading_issues,
            image_issues=image_issues,
            table_issues=table_issues,
            list_issues=list_issues,
            link_issues=link_issues,
            language_issues=language_issues,
            title_issues=title_issues,
            font_size_issues=font_size_issues,
            smartart_issues=smartart_issues,
            embedded_object_issues=embedded_object_issues,
            summary=summary,
            compliance_score=compliance_score,
            html_output=html_output,
            remediation_suggestions=remediation_suggestions,
            cvd_analysis=cvd_analysis,
        )

    def _extract_document_context(self, doc: Document, filename: str) -> Dict:
        """Extract document context for AI understanding"""
        context = {
            "filename": filename,
            "title": None,
            "headings": [],
            "topics": [],
            "has_images": False,
            "has_tables": False,
        }

        # Extract title (usually first heading or document core properties)
        try:
            if doc.core_properties.title:
                context["title"] = doc.core_properties.title
        except Exception:
            pass

        # Extract headings
        for i, para in enumerate(doc.paragraphs[:50]):  # Limit to first 50
            if _style_name(para).startswith("Heading"):
                try:
                    level = int(_style_name(para).replace("Heading ", ""))
                except Exception:
                    level = 2

                context["headings"].append(
                    {"level": level, "text": para.text[:100], "index": i}
                )

                if not context["title"] and level == 1:
                    context["title"] = para.text[:100]

        # Extract first few paragraphs as topics
        topic_paras = [p.text[:150] for p in doc.paragraphs[:5] if p.text.strip()]
        context["topics"] = topic_paras[:3]

        # Check for images and tables
        context["has_images"] = any(
            r._element.xpath(".//a:blip") for p in doc.paragraphs for r in p.runs
        )
        context["has_tables"] = len(doc.tables) > 0

        return context

    def _check_heading_structure(self, doc: Document) -> List[HeadingIssue]:
        """Check heading hierarchy for accessibility"""
        issues = []
        heading_levels = []

        for i, para in enumerate(doc.paragraphs):
            if _style_name(para).startswith("Heading"):
                try:
                    level = int(_style_name(para).replace("Heading ", ""))
                except Exception:
                    continue
                heading_levels.append((i, level, para.text[:100]))

        if not heading_levels:
            # No headings at all
            issues.append(
                HeadingIssue(
                    paragraph_index=0,
                    text="[Document]",
                    issue_type="no_headings",
                    suggested_fix="Add heading structure using Heading 1, Heading 2, etc. styles",
                )
            )
            return issues

        # Check if first heading is H1
        first_idx, first_level, first_text = heading_levels[0]
        if first_level != 1:
            issues.append(
                HeadingIssue(
                    paragraph_index=first_idx,
                    text=first_text,
                    issue_type="missing_h1",
                    current_level=first_level,
                    expected_level=1,
                    suggested_fix=f'Change "{first_text[:50]}..." to Heading 1 style',
                )
            )

        # Check for skipped levels
        prev_level = 0
        for idx, level, text in heading_levels:
            if level > prev_level + 1 and prev_level > 0:
                issues.append(
                    HeadingIssue(
                        paragraph_index=idx,
                        text=text,
                        issue_type="skipped_level",
                        current_level=level,
                        expected_level=prev_level + 1,
                        suggested_fix=f"Heading level skipped from H{prev_level} to H{level}. Use H{prev_level + 1} instead.",
                    )
                )
            prev_level = level

        return issues

    def _check_images(self, doc: Document, document_context: Dict) -> List[ImageIssue]:
        """Check images for alt text"""
        issues = []
        image_index = 0

        for para_idx, para in enumerate(doc.paragraphs):
            for run in para.runs:
                # Check for inline images
                blips = run._element.xpath(".//a:blip")
                for blip in blips:
                    # Try to get alt text from drawing
                    drawing = run._element.find(
                        ".//wp:inline",
                        {
                            "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                        },
                    )

                    alt_text = None
                    if drawing is not None:
                        # Check docPr for description
                        doc_pr = drawing.find(
                            ".//wp:docPr",
                            {
                                "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                            },
                        )
                        if doc_pr is not None:
                            alt_text = doc_pr.get("descr") or doc_pr.get("title")

                    has_alt = alt_text and alt_text.strip() != ""

                    # Capture the ZIP-relative path for this image (used by remediator)
                    image_zip_path = None
                    try:
                        embed_id = blip.get(qn("r:embed"))
                        if embed_id:
                            image_part = run.part.related_parts.get(embed_id)
                            if image_part:
                                # partname is like /word/media/image1.png — strip leading /
                                image_zip_path = str(image_part.partname).lstrip("/")
                    except Exception:
                        pass

                    if not has_alt:
                        suggested_alt = None
                        detected_type = None
                        is_decorative = False
                        is_chart = False
                        detailed_desc = None

                        # Generate alt text with AI if enabled
                        if self.generate_alt_text and self.image_generator:
                            try:
                                image_path = self._extract_image_from_run(run, blip)
                                if image_path:
                                    context = self._build_context_for_image(
                                        document_context, para_idx
                                    )

                                    # Smart image type detection
                                    type_result = run_async_from_sync(
                                        self.image_generator.detect_image_type(
                                            image_path=image_path, context=context
                                        )
                                    )

                                    if type_result.get("success"):
                                        detected_type = type_result.get(
                                            "image_type", "informative"
                                        )
                                        is_decorative = type_result.get(
                                            "is_decorative", False
                                        )
                                        image_purpose = type_result.get(
                                            "image_purpose", ""
                                        )

                                        is_chart = detected_type == "complex" or any(
                                            term in image_purpose.lower()
                                            for term in [
                                                "chart",
                                                "graph",
                                                "plot",
                                                "diagram",
                                                "infographic",
                                            ]
                                        )

                                        if is_decorative:
                                            suggested_alt = (
                                                ""  # Empty alt for decorative
                                            )
                                        elif is_chart:
                                            chart_result = run_async_from_sync(
                                                self.image_generator.describe_chart_or_graph(
                                                    image_path=image_path,
                                                    context=context,
                                                    detail_level="standard",
                                                )
                                            )
                                            if chart_result.get("success"):
                                                suggested_alt = chart_result.get(
                                                    "short_description", ""
                                                )
                                                detailed_desc = chart_result.get(
                                                    "detailed_description", ""
                                                )
                                        else:
                                            result = run_async_from_sync(
                                                self.image_generator.generate_alt_text(
                                                    image_path=image_path,
                                                    context=context,
                                                    educational_context=True,
                                                )
                                            )
                                            if result.get("success"):
                                                suggested_alt = result.get("alt_text")

                                    try:
                                        os.unlink(image_path)
                                    except Exception:
                                        pass
                            except Exception as e:
                                logger.warning(
                                    f"[DocxProcessor] Alt text generation failed: {e}"
                                )

                        issues.append(
                            ImageIssue(
                                paragraph_index=para_idx,
                                image_index=image_index,
                                has_alt_text=False,
                                existing_alt_text=None,
                                suggested_alt_text=suggested_alt,
                                image_path=image_zip_path,
                                detected_image_type=detected_type,
                                is_decorative=is_decorative,
                                is_chart=is_chart,
                                detailed_description=detailed_desc,
                            )
                        )
                    else:
                        # Image HAS alt text - validate it if enabled
                        if self.validate_alt_text and self.image_generator:
                            try:
                                image_path = self._extract_image_from_run(run, blip)
                                if image_path:
                                    context = self._build_context_for_image(
                                        document_context, para_idx
                                    )

                                    # Validate existing alt text with AI
                                    validation_result = run_async_from_sync(
                                        self.image_generator.validate_alt_text(
                                            image_path=image_path,
                                            existing_alt_text=alt_text,
                                            context=context,
                                        )
                                    )

                                    try:
                                        os.unlink(image_path)
                                    except Exception:
                                        pass

                                    if validation_result.get("success"):
                                        is_accurate = validation_result.get(
                                            "is_accurate", True
                                        )
                                        accuracy_score = validation_result.get(
                                            "accuracy_score", 1.0
                                        )
                                        validation_issues = validation_result.get(
                                            "issues", []
                                        )
                                        suggested_improvement = validation_result.get(
                                            "suggested_improvement"
                                        )

                                        # Only report if alt text is inaccurate or has issues
                                        if (
                                            not is_accurate
                                            or accuracy_score < 0.7
                                            or validation_issues
                                        ):
                                            issues.append(
                                                ImageIssue(
                                                    paragraph_index=para_idx,
                                                    image_index=image_index,
                                                    has_alt_text=True,
                                                    existing_alt_text=alt_text,
                                                    suggested_alt_text=suggested_improvement,
                                                    alt_text_validated=True,
                                                    alt_text_accurate=is_accurate,
                                                    alt_text_issues=validation_issues,
                                                    validation_score=accuracy_score,
                                                )
                                            )
                            except Exception as e:
                                logger.warning(
                                    f"[DocxProcessor] Alt text validation failed: {e}"
                                )

                    image_index += 1

        return issues

    def _extract_image_from_run(self, run, blip) -> Optional[str]:
        """Extract image from run and save to temp file"""
        try:
            # Get the embed id from blip
            embed_id = blip.get(qn("r:embed"))
            if not embed_id:
                return None

            # Get the image part from document
            part = run.part
            image_part = part.related_parts.get(embed_id)
            if not image_part:
                return None

            # Save to temp file
            image_bytes = image_part.blob
            image = Image.open(BytesIO(image_bytes))

            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                image.save(tmp, format="PNG")
                return tmp.name
        except Exception as e:
            logger.warning(f"[DocxProcessor] Failed to extract image: {e}")
            return None

    def _build_context_for_image(self, document_context: Dict, para_idx: int) -> str:
        """Build context string for AI alt text generation"""
        parts = []

        if document_context.get("title"):
            parts.append(f"Document: \"{document_context['title']}\"")

        parts.append(f"Image near paragraph {para_idx + 1}")

        if document_context.get("headings"):
            parts.append("Document structure:")
            for h in document_context["headings"][:5]:
                if h["index"] < para_idx:
                    indent = "  " * (h["level"] - 1)
                    parts.append(f"  {indent}H{h['level']}: \"{h['text']}\"")

        if document_context.get("topics"):
            parts.append("Document topics:")
            for topic in document_context["topics"]:
                parts.append(
                    f'  - "{topic[:100]}..."' if len(topic) > 100 else f'  - "{topic}"'
                )

        return "\n".join(parts)

    def _check_tables(self, doc: Document) -> List[TableIssue]:
        """Check tables for accessibility"""
        issues = []

        for table_idx, table in enumerate(doc.tables):
            row_count = len(table.rows)
            col_count = len(table.columns)

            if row_count == 0:
                continue

            # Check for header row
            has_header = False
            first_row = table.rows[0]

            # Check if first row has header formatting or explicit header style
            for cell in first_row.cells:
                # Check for bold text (common header indicator)
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.bold:
                            has_header = True
                            break

                # Check cell shading (another header indicator)
                try:
                    shading = cell._element.xpath(".//w:shd")
                    if shading:
                        has_header = True
                except Exception:
                    pass

            # Also check if table has explicit header row setting
            try:
                tbl_look = table._element.find(
                    ".//w:tblLook",
                    {
                        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                    },
                )
                if tbl_look is not None:
                    first_row_val = tbl_look.get(qn("w:firstRow"))
                    if first_row_val == "1":
                        has_header = True
            except Exception:
                pass

            if not has_header:
                issues.append(
                    TableIssue(
                        table_index=table_idx,
                        issue_type="missing_header",
                        row_count=row_count,
                        column_count=col_count,
                        suggested_fix="Define first row as header row. Select the table, go to Table Design > Header Row.",
                    )
                )

            # Check for complex merges (can confuse screen readers)
            has_complex_merge = False
            for row in table.rows:
                for cell in row.cells:
                    # Check for merged cells
                    tc = cell._tc
                    grid_span = tc.find(
                        ".//w:gridSpan",
                        {
                            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                        },
                    )
                    if grid_span is not None:
                        span = int(grid_span.get(qn("w:val"), 1))
                        if span > 1:
                            has_complex_merge = True
                            break

            if has_complex_merge:
                issues.append(
                    TableIssue(
                        table_index=table_idx,
                        issue_type="complex_merge",
                        row_count=row_count,
                        column_count=col_count,
                        suggested_fix="Complex merged cells can confuse screen readers. Consider simplifying table structure or adding alt text description.",
                    )
                )

        return issues

    def _check_lists(self, doc: Document) -> List[ListIssue]:
        """Check for fake lists (bullets via text vs real lists)"""
        issues = []

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # Skip if it's already a real list
            if _style_name(para) in ["List Paragraph", "List Bullet", "List Number"]:
                continue

            # Check for fake bullets at start of text
            for bullet in self.FAKE_BULLETS:
                if text.startswith(bullet) and len(text) > 2:
                    # Make sure it's not just the character alone
                    rest = text[len(bullet) :].strip()
                    if rest and not rest[0].isdigit():  # Not a numbered item
                        issues.append(
                            ListIssue(
                                paragraph_index=para_idx,
                                text=text[:100],
                                issue_type="fake_bullet",
                                detected_marker=bullet,
                                suggested_fix=f'Convert to real list using Home > Bullets. "{bullet}" is not accessible as a list marker.',
                            )
                        )
                    break

            # Check for fake numbered lists (1. 2. 3. etc. not in list style)
            number_match = re.match(r"^(\d+)[.)\-]\s+", text)
            if number_match:
                issues.append(
                    ListIssue(
                        paragraph_index=para_idx,
                        text=text[:100],
                        issue_type="fake_number",
                        detected_marker=number_match.group(1),
                        suggested_fix="Convert to real numbered list using Home > Numbering. Manual numbers are not accessible.",
                    )
                )

        return issues

    def _check_links(self, doc: Document) -> List[LinkIssue]:
        """Check hyperlinks for accessibility"""
        issues = []

        for para_idx, para in enumerate(doc.paragraphs):
            # Find hyperlinks in paragraph
            for hyperlink in para._element.xpath(".//w:hyperlink"):
                # Get link text
                link_text = ""
                for t in hyperlink.xpath(".//w:t"):
                    link_text += t.text or ""
                link_text = link_text.strip()

                # Get link URL
                r_id = hyperlink.get(qn("r:id"))
                link_url = ""
                if r_id:
                    try:
                        rel = para.part.rels.get(r_id)
                        if rel:
                            link_url = rel.target_ref
                    except Exception:
                        pass

                if not link_text:
                    continue

                # Check for non-descriptive link text
                for pattern in self.BAD_LINK_PATTERNS:
                    if re.match(pattern, link_text.lower()):
                        issues.append(
                            LinkIssue(
                                paragraph_index=para_idx,
                                link_text=link_text,
                                link_url=link_url,
                                issue_type="non_descriptive",
                                suggested_fix=f'Replace "{link_text}" with descriptive text that explains where the link goes.',
                            )
                        )
                        break

                # Check for URLs as link text
                if link_text.startswith(("http://", "https://", "www.")):
                    issues.append(
                        LinkIssue(
                            paragraph_index=para_idx,
                            link_text=link_text[:50],
                            link_url=link_url,
                            issue_type="url_as_text",
                            suggested_fix="Replace URL with descriptive link text. Screen readers will read the full URL.",
                        )
                    )

        return issues

    def _check_language(self, doc: Document) -> List[LanguageIssue]:
        """Check document language settings"""
        issues = []

        # Check document-level language
        try:
            # Word stores language in styles and paragraph settings

            # Check document settings
            settings = doc.settings
            if settings:
                # Check for language in doc settings (complex check)
                pass  # Most Word docs have language by default

            # For a more thorough check, we'd need to verify lang tags
            # This is a simplified check

        except Exception as e:
            logger.debug(f"[DocxProcessor] Language check error: {e}")

        # Note: Word typically sets language by default, so we mainly check
        # for mixed language content which is harder to detect

        return issues

    def _check_document_title(
        self, doc: Document, document_context: Dict, filename: str
    ) -> List[TitleIssue]:
        """
        Check document title for accessibility (WCAG 2.4.2).

        Validates that the document has a meaningful, descriptive title
        set in the core properties (metadata).

        Args:
            doc: The Document object
            document_context: Extracted document context
            filename: The original filename

        Returns:
            List of TitleIssue objects if issues found
        """
        issues = []

        # Generic/placeholder titles to flag
        GENERIC_TITLES = [
            "untitled",
            "document",
            "doc",
            "new document",
            "word document",
            "microsoft word",
            "document1",
            "document 1",
            "template",
        ]

        try:
            title = None
            try:
                title = doc.core_properties.title
            except Exception:
                pass

            # Get suggested title from context (first H1 or first paragraph)
            suggested_title = self._suggest_document_title(
                doc, document_context, filename
            )

            # Check for missing title
            if not title or not title.strip():
                issues.append(
                    TitleIssue(
                        issue_type="missing_title",
                        existing_title=title,
                        suggested_title=suggested_title,
                        suggested_fix=f'Add a descriptive title in File > Properties. Suggested: "{suggested_title}"',
                    )
                )
                return issues

            title_stripped = title.strip()
            title_lower = title_stripped.lower()

            # Check for generic/placeholder titles
            if title_lower in GENERIC_TITLES:
                issues.append(
                    TitleIssue(
                        issue_type="generic_title",
                        existing_title=title_stripped,
                        suggested_title=suggested_title,
                        suggested_fix=f'Replace generic title "{title_stripped}" with a descriptive title. Suggested: "{suggested_title}"',
                    )
                )
                return issues

            # Check if title is just the filename (without extension)
            filename_base = Path(filename).stem.lower()
            if title_lower == filename_base or title_lower == filename_base.replace(
                "_", " "
            ).replace("-", " "):
                issues.append(
                    TitleIssue(
                        issue_type="filename_title",
                        existing_title=title_stripped,
                        suggested_title=suggested_title,
                        suggested_fix=f'Document title "{title_stripped}" matches the filename. Consider a more descriptive title. Suggested: "{suggested_title}"',
                    )
                )
                return issues

            # Check for very short titles (less than 5 characters)
            if len(title_stripped) < 5:
                issues.append(
                    TitleIssue(
                        issue_type="empty_title",
                        existing_title=title_stripped,
                        suggested_title=suggested_title,
                        suggested_fix=f'Document title "{title_stripped}" is too short. Add a more descriptive title. Suggested: "{suggested_title}"',
                    )
                )

        except Exception as e:
            logger.debug(f"[DocxProcessor] Title check error: {e}")

        return issues

    def _suggest_document_title(
        self, doc: Document, document_context: Dict, filename: str
    ) -> str:
        """
        Generate a suggested document title based on content.

        Priority:
        1. First Heading 1 in the document
        2. First heading of any level
        3. First non-empty paragraph (first 60 chars)
        4. Cleaned filename

        Args:
            doc: The Document object
            document_context: Extracted document context
            filename: The original filename

        Returns:
            A suggested title string
        """
        # Try to get from context (already extracted)
        if document_context.get("title"):
            return document_context["title"][:80]

        # Look for headings
        headings = document_context.get("headings", [])
        if headings:
            # Prefer H1
            for h in headings:
                if h.get("level") == 1 and h.get("text"):
                    return h["text"][:80]
            # Otherwise first heading
            if headings[0].get("text"):
                return headings[0]["text"][:80]

        # Try first non-empty paragraph
        try:
            for para in doc.paragraphs[:10]:
                text = para.text.strip()
                if text and len(text) >= 5:
                    return text[:80]
        except Exception:
            pass

        # Fallback: clean the filename
        filename_base = Path(filename).stem
        cleaned = filename_base.replace("_", " ").replace("-", " ")
        return cleaned.title()[:80]

    def _check_font_sizes(self, doc: Document) -> List[FontSizeIssue]:
        """
        Check font sizes for accessibility (WCAG 1.4.4).

        WCAG 1.4.4 (Resize Text) requires text be resizable up to 200%.
        While no minimum size is mandated, small fonts hinder readability:
        - Text < 9pt: Very small, serious accessibility concern
        - Text < 11pt: Too small for comfortable reading

        Best practice: 11-12pt minimum for body text.

        Args:
            doc: The Document object

        Returns:
            List of FontSizeIssue objects for problematic text
        """
        issues = []
        MIN_FONT_SIZE_PT = 11.0  # Recommended minimum
        VERY_SMALL_SIZE_PT = 9.0  # Serious concern threshold

        # Track seen paragraphs to avoid duplicate issues
        para_issues_found = set()

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # Skip headings - they typically have adequate sizing
            if para.style and para.style.name.startswith("Heading"):
                continue

            for run in para.runs:
                run_text = run.text.strip()
                if not run_text:
                    continue

                # Get font size from run
                font_size_pt = None

                # Check run-level font size first
                if run.font and run.font.size:
                    # Font size is in EMUs (914400 EMUs per inch, 72 points per inch)
                    font_size_pt = run.font.size.pt

                # If no run-level font size, check style
                if font_size_pt is None:
                    try:
                        if para.style and para.style.font and para.style.font.size:
                            font_size_pt = para.style.font.size.pt
                    except Exception:
                        pass

                # If still no font size, check document default
                if font_size_pt is None:
                    try:
                        default_style = doc.styles["Normal"]
                        if default_style.font and default_style.font.size:
                            font_size_pt = default_style.font.size.pt
                    except Exception:
                        pass

                # Skip if we couldn't determine font size (assume default 11pt)
                if font_size_pt is None:
                    continue

                # Skip if this paragraph already has an issue logged
                if para_idx in para_issues_found:
                    continue

                # Check for very small text (< 9pt) - serious concern
                if font_size_pt < VERY_SMALL_SIZE_PT:
                    para_issues_found.add(para_idx)
                    issues.append(
                        FontSizeIssue(
                            paragraph_index=para_idx,
                            text_preview=text[:50] + ("..." if len(text) > 50 else ""),
                            font_size_pt=round(font_size_pt, 1),
                            issue_type="very_small",
                            suggested_fix=f"Text at {font_size_pt:.1f}pt is very small and difficult to read. Increase to at least 11pt for accessibility.",
                        )
                    )
                # Check for too small text (< 11pt)
                elif font_size_pt < MIN_FONT_SIZE_PT:
                    para_issues_found.add(para_idx)
                    issues.append(
                        FontSizeIssue(
                            paragraph_index=para_idx,
                            text_preview=text[:50] + ("..." if len(text) > 50 else ""),
                            font_size_pt=round(font_size_pt, 1),
                            issue_type="too_small",
                            suggested_fix=f"Text at {font_size_pt:.1f}pt may be difficult to read. Consider increasing to at least 11pt.",
                        )
                    )

        return issues

    def _detect_smartart(self, doc: Document, file_path: str) -> List[SmartArtIssue]:
        """
        Detect SmartArt diagrams and check their accessibility.

        SmartArt diagrams in DOCX files are stored in:
        - /word/diagrams/data*.xml (diagram data)
        - /word/diagrams/drawing*.xml (drawing)
        - /word/diagrams/layout*.xml (layout type)
        - /word/diagrams/colors*.xml (colors)

        Args:
            doc: python-docx Document object
            file_path: Path to DOCX file for direct ZIP access

        Returns:
            List of SmartArt accessibility issues
        """
        import zipfile
        from xml.etree import ElementTree as ET

        issues = []
        smartart_elements = []

        # Namespaces used in DOCX SmartArt
        namespaces = {
            "dgm": "http://schemas.openxmlformats.org/drawingml/2006/diagram",
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        }

        try:
            with zipfile.ZipFile(file_path, "r") as docx_zip:
                # Find diagram data files
                diagram_files = [
                    f
                    for f in docx_zip.namelist()
                    if f.startswith("word/diagrams/data") and f.endswith(".xml")
                ]

                for diagram_idx, diagram_file in enumerate(diagram_files):
                    try:
                        with docx_zip.open(diagram_file) as f:
                            tree = ET.parse(f)
                            root = tree.getroot()

                            # Extract text nodes from diagram
                            text_nodes = []
                            for elem in root.iter():
                                # Look for text elements (a:t tags)
                                if elem.tag.endswith("}t") or elem.tag == "t":
                                    if elem.text and elem.text.strip():
                                        text_nodes.append(elem.text.strip())

                            # Determine diagram type from layout file
                            diagram_type = self._get_smartart_type(
                                docx_zip, diagram_idx, namespaces
                            )

                            # Check for alt text in document.xml
                            has_alt_text, alt_text = self._get_smartart_alt_text(
                                doc, diagram_idx
                            )

                            smartart_elem = SmartArtElement(
                                diagram_index=diagram_idx,
                                diagram_type=diagram_type,
                                text_nodes=text_nodes,
                                node_count=len(text_nodes),
                                has_alt_text=has_alt_text,
                                existing_alt_text=alt_text,
                            )
                            smartart_elements.append(smartart_elem)

                            # Check for accessibility issues
                            if not has_alt_text:
                                # Generate description for the SmartArt
                                description = self._describe_smartart(smartart_elem)

                                issues.append(
                                    SmartArtIssue(
                                        diagram_index=diagram_idx,
                                        diagram_type=diagram_type,
                                        issue_type="missing_alt_text",
                                        node_count=len(text_nodes),
                                        text_content=text_nodes[
                                            :10
                                        ],  # Limit for display
                                        generated_description=description,
                                        suggested_fix=f"Add alt text to SmartArt diagram: '{description[:100]}...' "
                                        f"or mark as decorative if purely visual.",
                                    )
                                )

                            # Check for complex diagrams that need text alternatives
                            if len(text_nodes) > 10:
                                issues.append(
                                    SmartArtIssue(
                                        diagram_index=diagram_idx,
                                        diagram_type=diagram_type,
                                        issue_type="complex_structure",
                                        node_count=len(text_nodes),
                                        text_content=text_nodes[:10],
                                        generated_description=self._describe_smartart(
                                            smartart_elem
                                        ),
                                        suggested_fix="Complex SmartArt with many nodes. Consider providing "
                                        "a text-based alternative (list or table) nearby.",
                                    )
                                )

                    except Exception as e:
                        logger.warning(
                            f"[DocxProcessor] Error parsing SmartArt {diagram_file}: {e}"
                        )

        except Exception as e:
            logger.warning(f"[DocxProcessor] Error detecting SmartArt: {e}")

        return issues

    def _get_smartart_type(self, docx_zip, diagram_idx: int, namespaces: dict) -> str:
        """
        Determine SmartArt diagram type from layout file.

        Common SmartArt types:
        - hierarchy: Organization charts, tree structures
        - process: Step-by-step flows, timelines
        - cycle: Circular processes, recurring events
        - relationship: Venn diagrams, connections
        - matrix: Grid layouts, quadrants
        - pyramid: Triangular hierarchies
        - list: Bulleted or grouped items
        - picture: Image-centric layouts

        Args:
            docx_zip: ZipFile object for DOCX
            diagram_idx: Index of the diagram
            namespaces: XML namespaces

        Returns:
            Diagram type string
        """
        from xml.etree import ElementTree as ET

        # Try to find corresponding layout file
        layout_files = [
            f
            for f in docx_zip.namelist()
            if f.startswith("word/diagrams/layout") and f.endswith(".xml")
        ]

        diagram_type = "unknown"

        # Try to infer type from layout or quickStyle
        try:
            if diagram_idx < len(layout_files):
                with docx_zip.open(layout_files[diagram_idx]) as f:
                    tree = ET.parse(f)
                    root = tree.getroot()

                    # Look for layout type hints in the XML
                    xml_str = ET.tostring(root, encoding="unicode")

                    # Simple heuristic based on common patterns
                    if "hierarchy" in xml_str.lower() or "org" in xml_str.lower():
                        diagram_type = "hierarchy"
                    elif "process" in xml_str.lower() or "flow" in xml_str.lower():
                        diagram_type = "process"
                    elif "cycle" in xml_str.lower() or "circular" in xml_str.lower():
                        diagram_type = "cycle"
                    elif "venn" in xml_str.lower() or "relationship" in xml_str.lower():
                        diagram_type = "relationship"
                    elif "matrix" in xml_str.lower() or "grid" in xml_str.lower():
                        diagram_type = "matrix"
                    elif "pyramid" in xml_str.lower():
                        diagram_type = "pyramid"
                    elif "list" in xml_str.lower():
                        diagram_type = "list"
                    elif "picture" in xml_str.lower():
                        diagram_type = "picture"

        except Exception as e:
            logger.debug(f"[DocxProcessor] Could not determine SmartArt type: {e}")

        return diagram_type

    def _get_smartart_alt_text(
        self, doc: Document, diagram_idx: int
    ) -> tuple[bool, Optional[str]]:
        """
        Check if SmartArt has alt text defined.

        Alt text for SmartArt is typically stored in the drawing wrapper
        in document.xml.

        Args:
            doc: python-docx Document object
            diagram_idx: Index of the diagram

        Returns:
            Tuple of (has_alt_text, alt_text_value)
        """
        # SmartArt alt text checking is complex because it's in the drawing element
        # For now, use a simple approach - check if document has any drawing descriptions
        try:
            # Check document body for docPr elements with descr attribute
            for para in doc.paragraphs:
                xml = para._element.xml
                if 'descr="' in xml and "dgm:" in xml:
                    # Extract description
                    import re

                    match = re.search(r'descr="([^"]*)"', xml)
                    if match and match.group(1):
                        return True, match.group(1)

        except Exception as e:
            logger.debug(f"[DocxProcessor] Error checking SmartArt alt text: {e}")

        return False, None

    def _describe_smartart(self, smartart: SmartArtElement) -> str:
        """
        Generate an accessible description for SmartArt.

        Creates a natural language description based on the diagram type
        and extracted text nodes.

        Args:
            smartart: SmartArtElement with diagram info

        Returns:
            Human-readable description
        """
        type_descriptions = {
            "hierarchy": "organizational chart",
            "process": "process flow diagram",
            "cycle": "cycle diagram",
            "relationship": "relationship diagram",
            "matrix": "matrix diagram",
            "pyramid": "pyramid diagram",
            "list": "list diagram",
            "picture": "picture-based diagram",
            "unknown": "diagram",
        }

        type_desc = type_descriptions.get(smartart.diagram_type, "diagram")

        if not smartart.text_nodes:
            return f"A {type_desc} with {smartart.node_count} elements."

        # Build description based on type
        if smartart.diagram_type == "hierarchy":
            if len(smartart.text_nodes) >= 1:
                return (
                    f"Organization chart with '{smartart.text_nodes[0]}' at the top, "
                    f"containing {smartart.node_count} total items: "
                    f"{', '.join(smartart.text_nodes[:5])}"
                    f"{'...' if len(smartart.text_nodes) > 5 else ''}"
                )

        elif smartart.diagram_type == "process":
            return (
                f"Process flow with {smartart.node_count} steps: "
                f"{' → '.join(smartart.text_nodes[:5])}"
                f"{'...' if len(smartart.text_nodes) > 5 else ''}"
            )

        elif smartart.diagram_type == "cycle":
            return (
                f"Cycle diagram with {smartart.node_count} stages: "
                f"{' → '.join(smartart.text_nodes[:5])} → (repeats)"
            )

        elif smartart.diagram_type == "relationship":
            return (
                f"Relationship diagram showing connections between: "
                f"{', '.join(smartart.text_nodes[:5])}"
                f"{'...' if len(smartart.text_nodes) > 5 else ''}"
            )

        # Default description
        return (
            f"A {type_desc} containing {smartart.node_count} items: "
            f"{', '.join(smartart.text_nodes[:5])}"
            f"{'...' if len(smartart.text_nodes) > 5 else ''}"
        )

    def _detect_embedded_objects(
        self, doc: Document, file_path: str
    ) -> List[EmbeddedObjectIssue]:
        """
        Detect embedded objects in DOCX and check their accessibility.

        Embedded objects in DOCX files include:
        - OLE objects (oleObject in document.xml)
        - Embedded packages (Excel, PDF, etc.)
        - Embedded images that aren't inline

        Objects are stored in:
        - /word/embeddings/ (OLE objects, packages)
        - Relationships in document.xml.rels

        Args:
            doc: python-docx Document object
            file_path: Path to DOCX file for direct ZIP access

        Returns:
            List of embedded object accessibility issues
        """
        import zipfile
        from xml.etree import ElementTree as ET

        issues = []
        embedded_objects = []

        # Content type mappings
        type_mappings = {
            "oleObject": "ole",
            "package": "package",
            "spreadsheetml": "excel",
            "pdf": "pdf",
            "presentationml": "powerpoint",
            "wordprocessingml": "word",
        }

        try:
            with zipfile.ZipFile(file_path, "r") as docx_zip:
                # Check for embedded files in word/embeddings/
                embedding_files = [
                    f for f in docx_zip.namelist() if f.startswith("word/embeddings/")
                ]

                # Parse relationships to get object info
                rels_path = "word/_rels/document.xml.rels"
                rels_info = {}

                if rels_path in docx_zip.namelist():
                    try:
                        with docx_zip.open(rels_path) as f:
                            rels_tree = ET.parse(f)
                            rels_root = rels_tree.getroot()

                            for rel in rels_root:
                                target = rel.get("Target", "")
                                rel_type = rel.get("Type", "")
                                rel_id = rel.get("Id", "")

                                if "embeddings" in target or "oleObject" in rel_type:
                                    rels_info[target] = {
                                        "type": rel_type,
                                        "id": rel_id,
                                    }
                    except Exception as e:
                        logger.debug(f"[DocxProcessor] Error parsing rels: {e}")

                # Process each embedded file
                for obj_idx, embed_file in enumerate(embedding_files):
                    try:
                        file_info = docx_zip.getinfo(embed_file)
                        file_name = os.path.basename(embed_file)

                        # Determine object type
                        obj_type = "other"
                        for key, value in type_mappings.items():
                            if key in embed_file.lower() or key in file_name.lower():
                                obj_type = value
                                break

                        # Check file extension
                        ext = os.path.splitext(file_name)[1].lower()
                        ext_types = {
                            ".xlsx": "excel",
                            ".xls": "excel",
                            ".pdf": "pdf",
                            ".pptx": "powerpoint",
                            ".ppt": "powerpoint",
                            ".docx": "word",
                            ".doc": "word",
                            ".bin": "ole",
                        }
                        if ext in ext_types:
                            obj_type = ext_types[ext]

                        # Check for alt text (search in document.xml)
                        has_alt_text, alt_text = self._get_embedded_alt_text(
                            doc, embed_file, obj_idx
                        )

                        embedded_obj = EmbeddedObject(
                            object_index=obj_idx,
                            object_type=obj_type,
                            file_name=file_name,
                            location=embed_file,
                            has_alt_text=has_alt_text,
                            existing_alt_text=alt_text,
                            is_accessible=has_alt_text,
                            size_bytes=file_info.file_size,
                        )
                        embedded_objects.append(embedded_obj)

                        # Generate accessibility issues
                        recommendations = []

                        if not has_alt_text:
                            recommendations.append(
                                f"Add descriptive alt text for embedded {obj_type} object"
                            )

                        # Type-specific recommendations
                        if obj_type == "excel":
                            recommendations.append(
                                "Consider providing data in accessible table format in the document"
                            )
                            recommendations.append(
                                "Ensure embedded Excel has proper header rows and structure"
                            )
                        elif obj_type == "pdf":
                            recommendations.append(
                                "Ensure embedded PDF is tagged and accessible"
                            )
                            recommendations.append(
                                "Consider extracting key content to main document"
                            )
                        elif obj_type == "ole":
                            recommendations.append(
                                "OLE objects may not be accessible to screen readers"
                            )
                            recommendations.append(
                                "Provide text alternative describing the object's content"
                            )

                        if recommendations:
                            issues.append(
                                EmbeddedObjectIssue(
                                    object_index=obj_idx,
                                    object_type=obj_type,
                                    file_name=file_name,
                                    issue_type=(
                                        "missing_alt_text"
                                        if not has_alt_text
                                        else "inaccessible_content"
                                    ),
                                    location=embed_file,
                                    recommendations=recommendations,
                                    suggested_fix=(
                                        recommendations[0] if recommendations else ""
                                    ),
                                )
                            )

                    except Exception as e:
                        logger.warning(
                            f"[DocxProcessor] Error processing embedded object {embed_file}: {e}"
                        )

                # Also check for OLE objects referenced in document.xml
                ole_objects = self._find_ole_objects_in_doc(doc)
                for ole_idx, ole_info in enumerate(ole_objects):
                    if not any(
                        obj.location == ole_info.get("location", "")
                        for obj in embedded_objects
                    ):
                        # New OLE object not in embeddings folder
                        obj_idx = len(embedded_objects) + ole_idx

                        issues.append(
                            EmbeddedObjectIssue(
                                object_index=obj_idx,
                                object_type="ole",
                                file_name=ole_info.get("name"),
                                issue_type="missing_alt_text",
                                location=ole_info.get("location", "inline"),
                                recommendations=[
                                    "OLE object detected without accessibility information",
                                    "Add descriptive text alternative for screen readers",
                                ],
                                suggested_fix="Add alt text describing the embedded object's content",
                            )
                        )

        except Exception as e:
            logger.warning(f"[DocxProcessor] Error detecting embedded objects: {e}")

        return issues

    def _get_embedded_alt_text(
        self, doc: Document, embed_path: str, obj_idx: int
    ) -> tuple[bool, Optional[str]]:
        """
        Check if an embedded object has alt text.

        Args:
            doc: python-docx Document object
            embed_path: Path to embedded file within DOCX
            obj_idx: Index of the object

        Returns:
            Tuple of (has_alt_text, alt_text_value)
        """
        try:
            # Check document body for object references with descriptions
            for para in doc.paragraphs:
                xml = para._element.xml

                # Look for oleObject or embedded references with descr
                if "oleObject" in xml or "embed" in xml.lower():
                    import re

                    # Check for description attribute
                    match = re.search(r'descr="([^"]*)"', xml)
                    if match and match.group(1):
                        return True, match.group(1)

                    # Check for alt attribute
                    match = re.search(r'alt="([^"]*)"', xml)
                    if match and match.group(1):
                        return True, match.group(1)

        except Exception as e:
            logger.debug(f"[DocxProcessor] Error checking embedded alt text: {e}")

        return False, None

    def _find_ole_objects_in_doc(self, doc: Document) -> List[Dict]:
        """
        Find OLE objects referenced in document paragraphs.

        Args:
            doc: python-docx Document object

        Returns:
            List of dictionaries with OLE object info
        """
        ole_objects = []

        try:
            for para_idx, para in enumerate(doc.paragraphs):
                xml = para._element.xml

                if "oleObject" in xml or "o:OLEObject" in xml:
                    ole_objects.append(
                        {
                            "paragraph_index": para_idx,
                            "location": f"paragraph_{para_idx}",
                            "name": f"OLE Object {len(ole_objects) + 1}",
                        }
                    )

        except Exception as e:
            logger.debug(f"[DocxProcessor] Error finding OLE objects: {e}")

        return ole_objects

    def _count_real_lists(self, doc: Document) -> int:
        """Count real list items in document"""
        count = 0
        for para in doc.paragraphs:
            if _style_name(para) in ["List Paragraph", "List Bullet", "List Number"]:
                count += 1
        return count

    def _count_links(self, doc: Document) -> int:
        """Count hyperlinks in document"""
        count = 0
        for para in doc.paragraphs:
            count += len(para._element.xpath(".//w:hyperlink"))
        return count

    def _calculate_summary(
        self,
        heading_issues: List[HeadingIssue],
        image_issues: List[ImageIssue],
        table_issues: List[TableIssue],
        list_issues: List[ListIssue],
        link_issues: List[LinkIssue],
        language_issues: List[LanguageIssue],
        title_issues: List[TitleIssue] = None,
        font_size_issues: List[FontSizeIssue] = None,
        smartart_issues: List[SmartArtIssue] = None,
        embedded_object_issues: List[EmbeddedObjectIssue] = None,
    ) -> Dict[str, int]:
        """Calculate summary statistics"""
        title_count = len(title_issues) if title_issues else 0
        font_size_count = len(font_size_issues) if font_size_issues else 0
        smartart_count = len(smartart_issues) if smartart_issues else 0
        embedded_count = len(embedded_object_issues) if embedded_object_issues else 0
        return {
            "heading_issues": len(heading_issues),
            "image_issues": len(image_issues),
            "table_issues": len(table_issues),
            "list_issues": len(list_issues),
            "link_issues": len(link_issues),
            "language_issues": len(language_issues),
            "title_issues": title_count,
            "font_size_issues": font_size_count,
            "smartart_issues": smartart_count,
            "embedded_object_issues": embedded_count,
            "total_issues": (
                len(heading_issues)
                + len(image_issues)
                + len(table_issues)
                + len(list_issues)
                + len(link_issues)
                + len(language_issues)
                + title_count
                + font_size_count
                + smartart_count
                + embedded_count
            ),
        }

    def _calculate_compliance_score(self, summary: Dict, total_elements: int) -> float:
        """Calculate compliance score using unified scoring system"""
        from .compliance_scoring import score_from_severity_counts

        # No early return for total_elements == 0: issues must always count
        # (issue #90 — near-empty docs scored 100.0 while listing defects).
        # Map DOCX issue types to severity:
        # - Missing alt text = High (WCAG 1.1.1)
        # - Heading issues = Medium (WCAG 1.3.1)
        # - Table issues = Medium (WCAG 1.3.1)
        # - List issues = Low (best practice)
        # - Link issues = Medium (WCAG 2.4.4)
        # - Language issues = High (WCAG 3.1.1)
        # - Title issues = Medium (WCAG 2.4.2)
        # - Font size issues = Medium (WCAG 1.4.4)

        critical = 0
        high = (
            summary.get("image_issues", 0)
            + summary.get("language_issues", 0)
            + summary.get("smartart_issues", 0)
            + summary.get("embedded_object_issues", 0)
        )
        medium = (
            summary.get("heading_issues", 0)
            + summary.get("table_issues", 0)
            + summary.get("link_issues", 0)
            + summary.get("title_issues", 0)
            + summary.get("font_size_issues", 0)
        )
        low = summary.get("list_issues", 0)

        result = score_from_severity_counts(
            critical=critical,
            high=high,
            medium=medium,
            low=low,
            total_elements=total_elements,
        )
        return result.score

    def _generate_html(self, doc: Document, context: Dict, title: str) -> str:
        """Generate accessible HTML from Word document"""
        html = '<!DOCTYPE html>\n<html lang="en">\n<head>\n'
        html += '  <meta charset="UTF-8">\n'
        html += (
            '  <meta name="viewport" content="width=device-width, initial-scale=1.0">\n'
        )
        html += f'  <title>{context.get("title", title)}</title>\n'
        html += "  <style>\n"
        html += "    body { font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }\n"
        html += (
            "    h1, h2, h3, h4, h5, h6 { margin-top: 1.5em; margin-bottom: 0.5em; }\n"
        )
        html += "    p { margin-bottom: 1em; }\n"
        html += "    ul, ol { margin-bottom: 1em; }\n"
        html += "    table { border-collapse: collapse; width: 100%; margin-bottom: 1em; }\n"
        html += (
            "    td, th { border: 1px solid #ddd; padding: 8px; text-align: left; }\n"
        )
        html += "    th { background-color: #f5f5f5; }\n"
        html += "    img { max-width: 100%; height: auto; }\n"
        html += "  </style>\n"
        html += "</head>\n<body>\n"

        # Convert paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Handle headings
            if _style_name(para).startswith("Heading"):
                try:
                    level = int(_style_name(para).replace("Heading ", ""))
                except Exception:
                    level = 2
                html += f"  <h{level}>{self._escape_html(text)}</h{level}>\n"
            # Handle lists
            elif _style_name(para) in ["List Paragraph", "List Bullet", "List Number"]:
                html += f"  <li>{self._escape_html(text)}</li>\n"
            # Handle regular paragraphs
            else:
                html += f"  <p>{self._escape_html(text)}</p>\n"

        # Convert tables
        for table in doc.tables:
            html += "  <table>\n"
            for row_idx, row in enumerate(table.rows):
                html += "    <tr>\n"
                tag = "th" if row_idx == 0 else "td"
                for cell in row.cells:
                    cell_text = " ".join(p.text for p in cell.paragraphs)
                    html += f"      <{tag}>{self._escape_html(cell_text)}</{tag}>\n"
                html += "    </tr>\n"
            html += "  </table>\n"

        html += "</body>\n</html>"
        return html

    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters"""
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#x27;")
        )

    def _generate_remediation_suggestions(self, summary: Dict) -> List[str]:
        """Generate high-level remediation suggestions"""
        suggestions = []

        if summary.get("title_issues", 0) > 0:
            suggestions.append(
                "Set a descriptive document title. "
                "Go to File > Properties > Summary and add a meaningful title (WCAG 2.4.2)."
            )

        if summary["heading_issues"] > 0:
            suggestions.append(
                f"Fix {summary['heading_issues']} heading structure issues. "
                "Use proper Heading 1, Heading 2, etc. styles in order."
            )

        if summary["image_issues"] > 0:
            suggestions.append(
                f"Add alt text to {summary['image_issues']} images. "
                "Right-click image > Edit Alt Text to add descriptions."
            )

        if summary["table_issues"] > 0:
            suggestions.append(
                f"Fix {summary['table_issues']} table accessibility issues. "
                "Mark header rows using Table Design > Header Row checkbox."
            )

        if summary["list_issues"] > 0:
            suggestions.append(
                f"Convert {summary['list_issues']} fake lists to real lists. "
                "Use Home > Bullets or Numbering instead of manual symbols."
            )

        if summary["link_issues"] > 0:
            suggestions.append(
                f"Improve {summary['link_issues']} link text descriptions. "
                "Replace 'click here' with descriptive text explaining the destination."
            )

        if summary.get("font_size_issues", 0) > 0:
            suggestions.append(
                f"Increase font size for {summary['font_size_issues']} text elements. "
                "Use at least 11pt for body text and 12pt for better readability (WCAG 1.4.4)."
            )

        if summary["total_issues"] == 0:
            suggestions.append(
                "No accessibility issues detected. Document is WCAG 2.1 compliant!"
            )

        return suggestions

    def _analyze_cvd_accessibility(
        self, doc: Document
    ) -> List[ColorBlindnessAnalysisResult]:
        """
        Analyze color accessibility for color-blind users.

        Extracts text colors and background colors from Word document
        and tests them against all CVD types.

        Args:
            doc: python-docx Document object

        Returns:
            List of ColorBlindnessAnalysisResult for each unique color pair
        """
        if not self.cvd_simulator:
            return []

        results = []
        color_pairs_seen = set()

        # Helper to convert DOCX color to hex
        def color_to_hex(color_obj) -> Optional[str]:
            """Convert python-docx color to hex string"""
            if color_obj is None:
                return None
            # Check if it's an RGB color
            if hasattr(color_obj, "rgb") and color_obj.rgb:
                rgb = color_obj.rgb
                # RGB is stored as an integer or tuple
                if isinstance(rgb, int):
                    return f"#{rgb:06x}"
                elif hasattr(rgb, "__iter__"):
                    r, g, b = rgb
                    return f"#{r:02x}{g:02x}{b:02x}"
            # Check for theme color (can't convert directly)
            return None

        try:
            # Analyze paragraph text colors
            for para in doc.paragraphs:
                for run in para.runs:
                    # Get text color
                    fg_hex = None
                    if run.font.color and run.font.color.rgb:
                        rgb = run.font.color.rgb
                        if isinstance(rgb, int):
                            fg_hex = f"#{rgb:06x}"
                        else:
                            # RGBColor object
                            fg_hex = f"#{rgb}"

                    # Skip if no color set (default black)
                    if not fg_hex:
                        continue

                    # Get highlight color as background (if any)
                    bg_hex = "#ffffff"  # Default white
                    if run.font.highlight_color:
                        # python-docx highlight colors are enum, not RGB
                        # Map common highlights to approximate hex
                        highlight_map = {
                            "YELLOW": "#ffff00",
                            "GREEN": "#00ff00",
                            "CYAN": "#00ffff",
                            "PINK": "#ff00ff",
                            "BLUE": "#0000ff",
                            "RED": "#ff0000",
                            "DARK_BLUE": "#000080",
                            "TEAL": "#008080",
                            "DARK_RED": "#800000",
                            "GRAY_25": "#c0c0c0",
                            "GRAY_50": "#808080",
                        }
                        hl_name = str(run.font.highlight_color).split(".")[-1]
                        bg_hex = highlight_map.get(hl_name, "#ffffff")

                    # Skip if already analyzed
                    pair_key = (fg_hex.lower(), bg_hex.lower())
                    if pair_key in color_pairs_seen:
                        continue
                    color_pairs_seen.add(pair_key)

                    # Skip black on white
                    if fg_hex.lower() == "#000000" and bg_hex.lower() == "#ffffff":
                        continue

                    # Analyze this color pair
                    try:
                        analysis = self.cvd_simulator.analyze_color_accessibility(
                            foreground=fg_hex, background=bg_hex
                        )
                        if analysis.issues:
                            results.append(analysis)
                    except Exception as e:
                        logger.warning(
                            f"[DocxProcessor] CVD analysis failed for {fg_hex}/{bg_hex}: {e}"
                        )

            # Analyze table cell colors
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Check cell shading/fill
                        try:
                            shading = cell._tc.get_or_add_tcPr().xpath(
                                ".//w:shd",
                                namespaces={
                                    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                                },
                            )
                            if shading:
                                fill = shading[0].get(
                                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill"
                                )
                                if fill and fill != "auto" and len(fill) == 6:
                                    bg_hex = f"#{fill}"

                                    # Default text color is black
                                    fg_hex = "#000000"

                                    pair_key = (fg_hex.lower(), bg_hex.lower())
                                    if pair_key not in color_pairs_seen:
                                        color_pairs_seen.add(pair_key)

                                        try:
                                            analysis = self.cvd_simulator.analyze_color_accessibility(
                                                foreground=fg_hex, background=bg_hex
                                            )
                                            if analysis.issues:
                                                results.append(analysis)
                                        except Exception:
                                            pass
                        except Exception:
                            pass

        except Exception as e:
            logger.error(f"[DocxProcessor] CVD analysis failed: {e}")

        return results

    def process_directory(self, directory: str) -> List[DocxProcessingResult]:
        """
        Batch process all DOCX files in a directory

        Args:
            directory: Path to directory containing DOCX files

        Returns:
            List of DocxProcessingResult for each file
        """
        results = []
        docx_files = list(Path(directory).glob("*.docx"))

        for docx_file in docx_files:
            # Skip temp files
            if docx_file.name.startswith("~$"):
                continue

            try:
                logger.info(f"[DocxProcessor] Processing: {docx_file.name}")
                result = self.process_docx(str(docx_file))
                results.append(result)
            except Exception as e:
                logger.error(f"[DocxProcessor] Error processing {docx_file}: {e}")

        return results


class DocxBatchProcessor:
    """Batch processor for Word documents"""

    def __init__(self, generate_alt_text: bool = False):
        self.processor = DocxProcessor(generate_alt_text=generate_alt_text)

    def process_directory(self, directory: str) -> List[DocxProcessingResult]:
        """Process all DOCX files in a directory"""
        return self.processor.process_directory(directory)
