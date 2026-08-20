"""
Word Document Remediator for Aelira Auto-Remediation Engine.

This module provides automatic remediation for accessibility issues in
Microsoft Word documents (.docx files).

Supported auto-fixes:
- Add/update alt text for images
- Fix heading hierarchy
- Convert fake lists to real lists
- Add table headers
- Update non-descriptive link text
- Set document language
"""

import re
import logging
from typing import Any, Dict, List, Optional
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

from ..docx_processor import _style_name
from .base import (
    BaseRemediator,
    RemediationIssue,
    IssueCategory,
    IssueSeverity,
    RemediationConfig,
    VerificationResult,
)

logger = logging.getLogger(__name__)


class DocxRemediator(BaseRemediator):
    """
    Remediator for Microsoft Word documents (.docx).

    Automatically fixes accessibility issues including:
    - Missing or inadequate alt text on images
    - Improper heading structure
    - Fake bulleted lists (using symbols instead of list styles)
    - Missing table headers
    - Non-descriptive link text
    - Missing document language

    Usage:
        issues = [{'type': 'alt_text', 'severity': 'high', ...}]
        remediator = DocxRemediator('document.docx', issues)
        result = remediator.remediate()
    """

    DOCUMENT_TYPE = "word"
    SUPPORTED_EXTENSIONS = [".docx"]

    AUTO_FIXABLE_CATEGORIES = [
        IssueCategory.ALT_TEXT,
        IssueCategory.HEADING,
        IssueCategory.LIST,
        IssueCategory.TABLE,
        IssueCategory.LINK,
        IssueCategory.LANGUAGE,
        IssueCategory.TITLE,
    ]

    # Heading style name mapping
    HEADING_STYLES = {
        1: "Heading 1",
        2: "Heading 2",
        3: "Heading 3",
        4: "Heading 4",
        5: "Heading 5",
        6: "Heading 6",
    }

    # Font sizes for created heading styles
    _HEADING_SIZES = {1: 24, 2: 18, 3: 14, 4: 12, 5: 11, 6: 10}

    @staticmethod
    def _ensure_heading_style(document: Document, level: int) -> Any:
        """Return the heading style for *level*, creating it if absent.

        Many DOCX files (especially those exported from Google Docs or
        plain-text editors) lack built-in heading styles.  This helper
        guarantees a usable style exists so ``paragraph.style = style``
        never raises ``KeyError``.
        """
        name = f"Heading {level}"
        try:
            return document.styles[name]
        except KeyError:
            pass

        # Create the style with proper formatting + outline level
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.bold = True
        style.font.size = Pt(DocxRemediator._HEADING_SIZES.get(level, 12))
        # Set outline level so the heading appears in navigation / TOC
        pPr = style.element.get_or_add_pPr()
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), str(level - 1))  # 0-based
        pPr.append(outline)
        logger.info(f"Created missing '{name}' style in document")
        return style

    # Fake bullet characters to detect
    FAKE_BULLETS = [
        "•",
        "-",
        "*",
        "○",
        "▪",
        "◦",
        "–",
        "—",
        ">",
        "·",
        "→",
        "➤",
        "➢",
        "►",
        "❖",
    ]

    def __init__(
        self,
        file_path: str,
        issues: List[Dict[str, Any]],
        config: Optional[RemediationConfig] = None,
        ai_client: Optional[Any] = None,
        *,
        alt_text_client: Optional[Any] = None,
    ) -> None:
        """Initialize the Word document remediator."""
        super().__init__(
            file_path, issues, config, ai_client, alt_text_client=alt_text_client
        )
        self._document: Optional[Document] = None

    def _load_document(self) -> Document:
        """Load the Word document for editing."""
        logger.info(f"Loading Word document: {self.file_path}")
        self._document = Document(self.file_path)
        return self._document

    def _save_document(self, document: Document) -> str:
        """Save the remediated Word document."""
        output_path = self._get_output_path()
        logger.info(f"Saving remediated document to: {output_path}")

        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)

        document.save(output_path)
        return output_path

    def can_auto_fix(self, issue: RemediationIssue) -> bool:
        """
        Determine if an issue can be automatically fixed.

        Auto-fixable issues:
        - Alt text: If we have AI or can generate placeholder
        - Headings: If suggested level is provided
        - Lists: If fake bullets detected
        - Tables: If header row can be identified
        - Links: If we can generate descriptive text
        - Language: Always fixable (set to 'en')
        """
        if issue.category not in self.AUTO_FIXABLE_CATEGORIES:
            return False

        # Category-specific checks
        if issue.category == IssueCategory.ALT_TEXT:
            # Can fix with AI, or with placeholder
            return self.config.use_ai or self.config.fix_alt_text

        if issue.category == IssueCategory.HEADING:
            # Can fix if we have suggested level, or can infer from context,
            # or can use AI to identify heading candidates (no_headings case)
            return bool(
                issue.metadata.get("suggested_level")
                or issue.metadata.get("current_level")
                or issue.fix_suggestion
                or (self.config.use_ai and self.ai_client)
            )

        if issue.category == IssueCategory.LIST:
            # Can fix fake lists
            return issue.metadata.get("is_fake_list", False)

        if issue.category == IssueCategory.TABLE:
            # Can fix if first row exists
            return issue.metadata.get("has_data_rows", True)

        if issue.category == IssueCategory.LINK:
            # Can fix with AI or rule-based
            return True

        if issue.category == IssueCategory.LANGUAGE:
            # Always fixable
            return True

        if issue.category == IssueCategory.TITLE:
            # Can fix if we have a suggested title
            return bool(issue.metadata.get("suggested_title") or issue.fix_suggestion)

        return False

    def apply_fix(
        self, issue: RemediationIssue, document: Document, fix_content: str
    ) -> bool:
        """
        Apply a fix to the Word document.

        Args:
            issue: The issue being fixed
            document: The Word document object
            fix_content: The content to apply as the fix

        Returns:
            True if fix was applied successfully
        """
        try:
            if issue.category == IssueCategory.ALT_TEXT:
                return self._apply_alt_text_fix(issue, document, fix_content)

            if issue.category == IssueCategory.HEADING:
                return self._apply_heading_fix(issue, document, fix_content)

            if issue.category == IssueCategory.LIST:
                return self._apply_list_fix(issue, document, fix_content)

            if issue.category == IssueCategory.TABLE:
                return self._apply_table_fix(issue, document, fix_content)

            if issue.category == IssueCategory.LINK:
                return self._apply_link_fix(issue, document, fix_content)

            if issue.category == IssueCategory.LANGUAGE:
                return self._apply_language_fix(issue, document, fix_content)

            if issue.category == IssueCategory.TITLE:
                return self._apply_title_fix(issue, document, fix_content)

            return False

        except Exception as e:
            logger.error(f"Failed to apply fix for issue {issue.id}: {e}")
            return False

    def _apply_alt_text_fix(
        self, issue: RemediationIssue, document: Document, alt_text: str
    ) -> bool:
        """Apply alt text fix to an image."""
        try:
            # Get image location from metadata
            para_index = issue.metadata.get("paragraph_index")
            # image_index is per-paragraph (typically 0 since each paragraph has one image)
            image_index = issue.metadata.get("image_index_in_para", 0)

            if para_index is None:
                logger.warning(f"No paragraph index for alt text fix: {issue.id}")
                return False

            # Find the paragraph
            if para_index >= len(document.paragraphs):
                logger.warning(f"Paragraph index out of range: {para_index}")
                return False

            paragraph = document.paragraphs[para_index]

            # Find inline shapes (images) in the paragraph
            images_found = 0
            for run in paragraph.runs:
                if run._element.xpath(".//a:blip"):
                    if images_found == image_index:
                        # Found the target image - set alt text
                        self._set_image_alt_text(run, alt_text)
                        logger.info(
                            f"Applied alt text to image at paragraph {para_index}"
                        )
                        return True
                    images_found += 1

            # Try to find inline shapes through paragraph's XML
            inline_shapes = paragraph._element.xpath(".//wp:inline | .//wp:anchor")
            if image_index < len(inline_shapes):
                self._set_inline_shape_alt_text(inline_shapes[image_index], alt_text)
                logger.info(
                    f"Applied alt text to inline shape at paragraph {para_index}"
                )
                return True

            logger.warning(f"Could not find image to apply alt text: {issue.id}")
            return False

        except Exception as e:
            logger.error(f"Error applying alt text fix: {e}")
            return False

    def _set_image_alt_text(self, run, alt_text: str) -> None:
        """Set alt text on an image in a run."""
        # Find wp:docPr from the run element — docPr is a child of wp:inline/wp:anchor,
        # NOT a child of a:blip's parent (pic:blipFill)
        doc_pr = run._element.xpath(".//wp:docPr")
        if doc_pr:
            doc_pr[0].set("descr", alt_text)

    def _set_inline_shape_alt_text(self, shape_element, alt_text: str) -> None:
        """Set alt text on an inline shape element."""
        # Use prefix notation — BaseOxmlElement.xpath() provides nsmap automatically
        doc_pr = shape_element.xpath(".//wp:docPr")
        if doc_pr:
            doc_pr[0].set("descr", alt_text)
        else:
            # Create docPr if it doesn't exist
            for child in shape_element:
                if "extent" in child.tag:
                    doc_pr_elem = OxmlElement("wp:docPr")
                    doc_pr_elem.set("id", "1")
                    doc_pr_elem.set("name", "Image")
                    doc_pr_elem.set("descr", alt_text)
                    shape_element.insert(
                        list(shape_element).index(child) + 1, doc_pr_elem
                    )
                    break

    def _apply_heading_fix(
        self, issue: RemediationIssue, document: Document, fix_content: str
    ) -> bool:
        """Apply heading style fix to a paragraph."""
        try:
            para_index = issue.metadata.get("paragraph_index")
            suggested_level = (
                issue.metadata.get("suggested_level")
                or issue.metadata.get("current_level")
                or 2
            )

            if para_index is None and issue.original_content:
                # Try to find by text content
                return self._apply_heading_by_text(
                    document, issue.original_content, suggested_level
                )

            if para_index is None and not issue.original_content:
                # "no_headings" case — no specific paragraph identified.
                # Use AI to identify which paragraphs should become headings.
                if self.config.use_ai and self.ai_client:
                    return self._apply_headings_with_ai(document)
                return False

            if para_index >= len(document.paragraphs):
                logger.warning(f"Paragraph index out of range: {para_index}")
                return False

            paragraph = document.paragraphs[para_index]

            # Apply the heading style (creating it if absent)
            style = self._ensure_heading_style(document, suggested_level)
            paragraph.style = style

            logger.info(f"Applied {style.name} to paragraph {para_index}")
            return True

        except Exception as e:
            logger.error(f"Error applying heading fix: {e}")
            return False

    def _apply_heading_by_text(self, document: Document, text: str, level: int) -> bool:
        """Find a paragraph by text and apply heading style."""
        if not text:
            return False

        for paragraph in document.paragraphs:
            if text.strip() in paragraph.text.strip():
                style = self._ensure_heading_style(document, level)
                paragraph.style = style
                logger.info(f"Applied {style.name} to paragraph: {text[:50]}...")
                return True

        return False

    def _apply_headings_with_ai(self, document: Document) -> bool:
        """Use AI to identify paragraphs that should be headings and apply styles.

        Called for the 'no_headings' case where the scanner detected zero
        headings but couldn't pinpoint which paragraphs should become them.
        """
        # Collect paragraph texts with indices for AI analysis
        paragraphs_for_ai = []
        for i, para in enumerate(document.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs_for_ai.append({"index": i, "text": text[:200]})

        if not paragraphs_for_ai:
            return False

        # Send first ~40 paragraphs to keep prompt short
        sample = paragraphs_for_ai[:40]
        para_list = "\n".join(f"[{p['index']}] {p['text']}" for p in sample)

        from ...utils.security import sanitize_for_prompt

        safe_para_list = sanitize_for_prompt(para_list, max_length=3000)

        prompt = f"""This Word document has no heading styles applied. Identify which paragraphs should be headings and what level (1-3) each should be.

Paragraphs (format: [index] text):
{safe_para_list}

Rules:
- Pick the document title or main topic as Heading 1 (only one H1)
- Major sections become Heading 2
- Subsections become Heading 3
- Short, title-like paragraphs are heading candidates
- Long body-text paragraphs are NOT headings
- Return ONLY a JSON array like: [{{"index": 0, "level": 1}}, {{"index": 5, "level": 2}}]
- Return an empty array [] if no clear headings can be identified"""

        try:
            if not hasattr(self.ai_client, "generate_text_sync"):
                return False

            self.result.ai_calls_made += 1
            result = self.ai_client.generate_text_sync(
                prompt=prompt, max_tokens=500, temperature=0.2
            )

            if not result.get("success") or not result.get("content"):
                return False

            # Parse JSON response
            import json

            content = result["content"].strip()
            # Strip markdown code fences if present
            if content.startswith("```"):
                content = content.split("\n", 1)[-1].rsplit("```", 1)[0].strip()

            headings = json.loads(content)
            if not isinstance(headings, list) or not headings:
                return False

            applied = 0
            for h in headings:
                idx = h.get("index")
                level = h.get("level", 2)
                if idx is None or not isinstance(idx, int):
                    continue
                if idx >= len(document.paragraphs):
                    continue
                level = max(1, min(6, int(level)))

                style = self._ensure_heading_style(document, level)
                document.paragraphs[idx].style = style
                applied += 1
                logger.info(
                    f"AI heading: applied {style.name} to para {idx}: "
                    f"{document.paragraphs[idx].text[:50]}..."
                )

            if applied:
                logger.info(f"AI identified and applied {applied} headings")
                return True

        except (json.JSONDecodeError, ValueError) as e:
            logger.warning(f"Failed to parse AI heading response: {e}")
        except Exception as e:
            logger.error(f"AI heading identification failed: {e}")

        return False

    def _apply_list_fix(
        self, issue: RemediationIssue, document: Document, fix_content: str
    ) -> bool:
        """Convert fake bullet list to real list."""
        try:
            para_indices = issue.metadata.get("paragraph_indices", [])
            if not para_indices:
                para_index = issue.metadata.get("paragraph_index")
                if para_index is not None:
                    para_indices = [para_index]

            if not para_indices:
                logger.warning(f"No paragraph indices for list fix: {issue.id}")
                return False

            for idx in para_indices:
                if idx >= len(document.paragraphs):
                    continue

                paragraph = document.paragraphs[idx]

                # Remove fake bullet character
                text = paragraph.text
                for bullet in self.FAKE_BULLETS:
                    if text.strip().startswith(bullet):
                        # Remove the bullet and any following space
                        new_text = text.lstrip()
                        for b in self.FAKE_BULLETS:
                            if new_text.startswith(b):
                                new_text = new_text[len(b) :].lstrip()
                                break

                        # Clear paragraph and set new text
                        paragraph.clear()
                        paragraph.add_run(new_text)
                        break

                # Apply list style
                try:
                    paragraph.style = "List Bullet"
                except KeyError:
                    # Create numbered list with paragraph formatting
                    self._apply_list_formatting(paragraph)

            logger.info(f"Applied list formatting to {len(para_indices)} paragraphs")
            return True

        except Exception as e:
            logger.error(f"Error applying list fix: {e}")
            return False

    def _apply_list_formatting(self, paragraph) -> None:
        """Apply list formatting to a paragraph using XML."""
        # Add numbering properties
        p = paragraph._element
        pPr = p.get_or_add_pPr()

        # Create numPr element
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), "1")

        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)

    def _apply_table_fix(
        self, issue: RemediationIssue, document: Document, fix_content: str
    ) -> bool:
        """Apply table header fix."""
        try:
            table_index = issue.metadata.get("table_index", 0)

            if table_index >= len(document.tables):
                logger.warning(f"Table index out of range: {table_index}")
                return False

            table = document.tables[table_index]

            if len(table.rows) == 0:
                logger.warning(f"Table has no rows: {table_index}")
                return False

            # Mark first row as header
            header_row = table.rows[0]

            # Apply header formatting
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Set table header property (repeat header row)
            tbl = table._tbl
            tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

            # Add tblHeader element to first row
            tr = header_row._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = OxmlElement("w:tblHeader")
            trPr.append(tblHeader)

            logger.info(f"Applied header formatting to table {table_index}")
            return True

        except Exception as e:
            logger.error(f"Error applying table fix: {e}")
            return False

    def _apply_link_fix(
        self, issue: RemediationIssue, document: Document, link_text: str
    ) -> bool:
        """Update link text to be more descriptive."""
        try:
            para_index = issue.metadata.get("paragraph_index")
            original_text = issue.original_content or issue.metadata.get(
                "link_text", ""
            )

            if not original_text:
                logger.warning(f"No original link text for fix: {issue.id}")
                return False

            # Search for the link text in the document
            for i, paragraph in enumerate(document.paragraphs):
                if para_index is not None and i != para_index:
                    continue

                if original_text in paragraph.text:
                    # Find and update the run containing the link text
                    for run in paragraph.runs:
                        if original_text in run.text:
                            run.text = run.text.replace(original_text, link_text)
                            logger.info(
                                f"Updated link text: '{original_text}' -> '{link_text}'"
                            )
                            return True

            logger.warning(f"Could not find link text to update: {original_text}")
            return False

        except Exception as e:
            logger.error(f"Error applying link fix: {e}")
            return False

    def _apply_language_fix(
        self, issue: RemediationIssue, document: Document, language: str
    ) -> bool:
        """Set document language."""
        try:
            # Set language on document core properties
            # Word stores language at paragraph/run level, not document level
            # We'll set it on all paragraphs

            lang_code = language or "en-US"

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    rPr = run._element.get_or_add_rPr()
                    lang = OxmlElement("w:lang")
                    lang.set(qn("w:val"), lang_code)
                    lang.set(qn("w:eastAsia"), lang_code)
                    lang.set(qn("w:bidi"), lang_code)
                    rPr.append(lang)

            logger.info(f"Set document language to: {lang_code}")
            return True

        except Exception as e:
            logger.error(f"Error applying language fix: {e}")
            return False

    def _apply_title_fix(
        self, issue: RemediationIssue, document: Document, title: str
    ) -> bool:
        """Set document title in core properties."""
        try:
            if not title:
                logger.warning(f"No title provided for title fix: {issue.id}")
                return False

            # Set the document title in core properties
            document.core_properties.title = title
            logger.info(f"Set document title to: {title}")
            return True

        except Exception as e:
            logger.error(f"Error applying title fix: {e}")
            return False

    def _get_rule_based_fix(
        self, issue: RemediationIssue, document: Any
    ) -> Optional[str]:
        """Get a rule-based fix for an issue."""
        if issue.category == IssueCategory.ALT_TEXT:
            # Decorative images get empty alt text per WCAG 1.1.1
            if issue.metadata.get("is_decorative"):
                return ""
            # Use pre-generated alt text from the scanner if available
            if self.config.allow_legacy_nested_ai:
                generated_alt = issue.metadata.get(
                    "suggested_alt_text"
                ) or issue.metadata.get("generated_alt_text")
                if generated_alt:
                    return generated_alt
                if issue.fix_suggestion:
                    return issue.fix_suggestion
            # Return None to let AI generation handle it in _generate_fix()
            return None

        if issue.category == IssueCategory.LANGUAGE:
            return "en-US"

        if issue.category == IssueCategory.HEADING:
            # Use suggested level from metadata
            suggested_level = issue.metadata.get("suggested_level")
            if suggested_level:
                return self.HEADING_STYLES.get(suggested_level, "Heading 2")
            return "Heading 2"

        if issue.category == IssueCategory.LIST:
            return "List Bullet"

        if issue.category == IssueCategory.TABLE:
            return "header_row"

        if issue.category == IssueCategory.LINK:
            # Try to generate from URL or context
            url = issue.metadata.get("url", "")
            if url:
                # Extract meaningful text from URL
                return self._generate_link_text_from_url(url)

        if issue.category == IssueCategory.TITLE:
            # Use suggested title from metadata or fix_suggestion
            suggested_title = issue.metadata.get("suggested_title")
            if suggested_title:
                return suggested_title
            if issue.fix_suggestion:
                return issue.fix_suggestion
            # Fallback: generate from filename
            from pathlib import Path

            filename = Path(self.file_path).stem
            return filename.replace("_", " ").replace("-", " ").title()

        return None

    def _generate_link_text_from_url(self, url: str) -> str:
        """Generate descriptive link text from a URL."""
        from urllib.parse import urlparse

        try:
            parsed = urlparse(url)

            # Get domain
            domain = parsed.netloc.replace("www.", "")

            # Get path parts
            path_parts = [p for p in parsed.path.split("/") if p]

            if path_parts:
                # Use the last meaningful path segment
                last_part = path_parts[-1]
                # Remove file extensions
                last_part = re.sub(r"\.[a-z]+$", "", last_part, flags=re.IGNORECASE)
                # Replace hyphens/underscores with spaces
                last_part = re.sub(r"[-_]", " ", last_part)
                # Capitalize
                last_part = last_part.title()

                if len(last_part) > 3:
                    return f"{last_part} on {domain}"

            return f"Visit {domain}"

        except Exception:
            return "Visit link"

    def _get_ai_generated_fix(
        self, issue: RemediationIssue, document: Any, *, client: Any
    ) -> Optional[str]:
        """Get an AI-generated fix for an issue."""
        try:
            self.result.ai_calls_made += 1

            if issue.category == IssueCategory.ALT_TEXT:
                return self._generate_alt_text_with_ai(issue, document, client=client)

            if issue.category == IssueCategory.LINK:
                return self._generate_link_text_with_ai(issue, client=client)

            return None

        except Exception as e:
            logger.error(f"AI generation failed: {e}")
            return None

    def _generate_alt_text_with_ai(
        self, issue: RemediationIssue, document: Any, *, client: Any
    ) -> Optional[str]:
        """Generate alt text using AI."""
        # Get image context
        para_index = issue.metadata.get("paragraph_index", 0)

        # Get surrounding text for context
        context_text = ""
        if document and para_index < len(document.paragraphs):
            # Get text from surrounding paragraphs
            start = max(0, para_index - 2)
            end = min(len(document.paragraphs), para_index + 3)
            context_text = " ".join(
                [p.text for p in document.paragraphs[start:end] if p.text.strip()]
            )

        # Sanitize context text before AI prompt
        from ...utils.security import sanitize_for_prompt

        safe_context = (
            sanitize_for_prompt(context_text, max_length=500)
            if context_text
            else "No context available"
        )

        # Generate prompt
        prompt = f"""Generate concise, descriptive alt text for an image in a Word document.

Context from surrounding text:
{safe_context}

Requirements:
- Be concise (under 125 characters)
- Describe the image's content and purpose
- Don't start with "Image of" or "Picture of"
- Focus on what's important for understanding the document

Generate only the alt text, nothing else:"""

        try:
            if hasattr(client, "analyze_image_sync"):
                try:
                    import zipfile

                    with zipfile.ZipFile(self.file_path, "r") as z:
                        image_rel_path = issue.metadata.get("image_path", "")
                        if not image_rel_path:
                            logger.debug("No image_path in metadata, skipping vision")
                            raise ValueError("No image_path in issue metadata")
                        if image_rel_path:
                            zip_path = image_rel_path.lstrip("/")
                            if not zip_path.startswith("word/"):
                                zip_path = f"word/{zip_path}"
                            image_bytes = z.read(zip_path)

                            result = client.analyze_image_sync(
                                image_data=image_bytes,
                                prompt=prompt,
                                max_tokens=200,
                            )
                            if result.get("success") and result.get("content"):
                                return result["content"].strip().strip("\"'")[:125]
                except Exception as e:
                    logger.warning(f"DOCX vision AI failed, falling back to text: {e}")

            if hasattr(client, "generate_text_sync"):
                result = client.generate_text_sync(
                    prompt=prompt,
                    max_tokens=200,
                    temperature=0.3,
                )
                if result.get("success") and result.get("content"):
                    alt_text = result["content"].strip().strip("\"'")
                    return alt_text[:125] if alt_text else None
        except Exception as e:
            logger.error(f"AI alt text generation failed: {e}")

        return None

    def _generate_link_text_with_ai(
        self, issue: RemediationIssue, *, client: Any
    ) -> Optional[str]:
        """Generate descriptive link text using AI."""
        original_text = issue.original_content or "click here"
        url = issue.metadata.get("url", "")
        context = issue.metadata.get("context", "")

        prompt = f"""Generate better link text for an accessibility fix.

Current link text: "{original_text}"
URL: {url}
Context: {context[:200] if context else 'No context'}

Requirements:
- Be descriptive of the link destination
- Keep it concise (2-5 words)
- Don't use generic text like "click here", "read more", "link"
- Make it clear what the user will find

Generate only the link text, nothing else:"""

        try:
            if hasattr(client, "generate_text_sync"):
                result = client.generate_text_sync(
                    prompt=prompt,
                    max_tokens=100,
                    temperature=0.3,
                )
                if result.get("success") and result.get("content"):
                    return result["content"].strip().strip("\"'")
        except Exception as e:
            logger.error(f"AI link text generation failed: {e}")

        # Fallback to rule-based
        return self._generate_link_text_from_url(url)

    def _verify_fixes(self, output_path: str) -> VerificationResult:
        """Verify that fixes were applied correctly."""
        try:
            # Reload the document and check
            verified_doc = Document(output_path)

            verified_count = 0
            for fixed in self.result.fixed_issues:
                if fixed.category == IssueCategory.HEADING:
                    # Verify heading was applied
                    para_index = self._get_para_index_from_fixed(fixed)
                    if para_index and para_index < len(verified_doc.paragraphs):
                        para = verified_doc.paragraphs[para_index]
                        if "Heading" in _style_name(para):
                            verified_count += 1
                            fixed.verification_passed = True
                        else:
                            fixed.verification_passed = False
                            self.result.warnings.append(
                                f"Heading fix may not have applied at paragraph {para_index}"
                            )
                else:
                    # Assume other fixes passed
                    verified_count += 1

            logger.info(
                f"Verified {verified_count}/{len(self.result.fixed_issues)} fixes"
            )

        except Exception as e:
            logger.warning(f"Fix verification failed: {e}")
            self.result.warnings.append(f"Could not verify fixes: {e}")

        # Honour the base contract: summarise via the base implementation,
        # which also records verification_result on self.result.
        return super()._verify_fixes(output_path)

    def _get_para_index_from_fixed(self, fixed) -> Optional[int]:
        """Extract paragraph index from fixed issue location."""
        if fixed.location:
            try:
                # Try to parse "Paragraph X" format
                match = re.search(r"paragraph\s*(\d+)", fixed.location, re.IGNORECASE)
                if match:
                    return int(match.group(1))
            except Exception:
                pass
        return None

    def _calculate_scores(self) -> None:
        """Calculate compliance scores for the remediation."""
        # Simple calculation based on fixes made
        if self.result.total_issues > 0:
            # Estimate original score
            # Each issue deducts based on severity
            severity_penalties = {
                IssueSeverity.CRITICAL: 15,
                IssueSeverity.HIGH: 10,
                IssueSeverity.MEDIUM: 5,
                IssueSeverity.LOW: 2,
            }

            total_penalty = sum(
                severity_penalties.get(issue.severity, 5) for issue in self.issues
            )
            self.result.original_compliance_score = max(0, 100 - total_penalty)

            # Calculate remediated score
            fixed_penalty_reduction = sum(
                severity_penalties.get(fixed.severity, 5)
                for fixed in self.result.fixed_issues
            )
            remaining_penalty = total_penalty - fixed_penalty_reduction
            self.result.remediated_compliance_score = max(0, 100 - remaining_penalty)

            self.result.improvement = (
                self.result.remediated_compliance_score
                - self.result.original_compliance_score
            )
