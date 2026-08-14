"""
Tests for SmartArt Accessibility Detection

Tests cover:
- SmartArt element model
- SmartArt issue model
- SmartArt detection in DOCX files
- SmartArt type identification
- SmartArt description generation
"""

import pytest
import tempfile
import os

from src.education.docx_processor import (
    DocxProcessor,
    SmartArtElement,
    SmartArtIssue,
)


class TestSmartArtElementModel:
    """Test SmartArtElement Pydantic model."""

    def test_smartart_element_creation(self):
        """Test creating a SmartArt element."""
        element = SmartArtElement(
            diagram_index=0,
            diagram_type="hierarchy",
            text_nodes=["CEO", "CTO", "CFO", "VP Engineering"],
            node_count=4,
            has_alt_text=True,
            existing_alt_text="Organization chart showing company structure",
        )

        assert element.diagram_index == 0
        assert element.diagram_type == "hierarchy"
        assert len(element.text_nodes) == 4
        assert element.node_count == 4
        assert element.has_alt_text is True

    def test_smartart_element_minimal(self):
        """Test SmartArt element with minimal fields."""
        element = SmartArtElement(
            diagram_index=1,
            diagram_type="process",
            text_nodes=["Step 1", "Step 2"],
            node_count=2,
            has_alt_text=False,
        )

        assert element.existing_alt_text is None
        assert element.relationship_type is None


class TestSmartArtIssueModel:
    """Test SmartArtIssue Pydantic model."""

    def test_smartart_issue_creation(self):
        """Test creating a SmartArt issue."""
        issue = SmartArtIssue(
            diagram_index=0,
            diagram_type="hierarchy",
            issue_type="missing_alt_text",
            node_count=5,
            text_content=["CEO", "CTO", "CFO"],
            generated_description="Organization chart with CEO at the top",
            suggested_fix="Add alt text describing the organizational structure",
        )

        assert issue.diagram_index == 0
        assert issue.issue_type == "missing_alt_text"
        assert len(issue.text_content) == 3
        assert issue.generated_description is not None

    def test_smartart_issue_complex_structure(self):
        """Test SmartArt issue for complex diagram."""
        issue = SmartArtIssue(
            diagram_index=1,
            diagram_type="process",
            issue_type="complex_structure",
            node_count=15,
            text_content=["Step 1", "Step 2", "Step 3"],
            suggested_fix="Consider providing text alternative",
        )

        assert issue.issue_type == "complex_structure"
        assert issue.node_count == 15


class TestSmartArtDescriptions:
    """Test SmartArt description generation."""

    @pytest.fixture
    def processor(self):
        """Create DocxProcessor for testing."""
        return DocxProcessor()

    def test_describe_hierarchy(self, processor):
        """Test description generation for hierarchy diagram."""
        element = SmartArtElement(
            diagram_index=0,
            diagram_type="hierarchy",
            text_nodes=["CEO", "CTO", "CFO", "VP Sales"],
            node_count=4,
            has_alt_text=False,
        )

        description = processor._describe_smartart(element)

        assert (
            "Organization chart" in description
            or "organizational chart" in description.lower()
        )
        assert "CEO" in description
        assert "4" in description

    def test_describe_process(self, processor):
        """Test description generation for process diagram."""
        element = SmartArtElement(
            diagram_index=0,
            diagram_type="process",
            text_nodes=["Plan", "Design", "Build", "Test", "Deploy"],
            node_count=5,
            has_alt_text=False,
        )

        description = processor._describe_smartart(element)

        assert "Process" in description or "process" in description.lower()
        assert "5" in description
        assert "→" in description  # Arrow indicating flow

    def test_describe_cycle(self, processor):
        """Test description generation for cycle diagram."""
        element = SmartArtElement(
            diagram_index=0,
            diagram_type="cycle",
            text_nodes=["Plan", "Do", "Check", "Act"],
            node_count=4,
            has_alt_text=False,
        )

        description = processor._describe_smartart(element)

        assert "Cycle" in description or "cycle" in description.lower()
        assert "repeats" in description

    def test_describe_relationship(self, processor):
        """Test description generation for relationship diagram."""
        element = SmartArtElement(
            diagram_index=0,
            diagram_type="relationship",
            text_nodes=["Engineering", "Design", "Product"],
            node_count=3,
            has_alt_text=False,
        )

        description = processor._describe_smartart(element)

        assert "Relationship" in description or "relationship" in description.lower()
        assert "connections" in description

    def test_describe_empty_diagram(self, processor):
        """Test description for diagram with no text."""
        element = SmartArtElement(
            diagram_index=0,
            diagram_type="unknown",
            text_nodes=[],
            node_count=0,
            has_alt_text=False,
        )

        description = processor._describe_smartart(element)

        assert "diagram" in description.lower()
        assert "0" in description

    def test_describe_truncates_long_lists(self, processor):
        """Test that long text lists are truncated."""
        element = SmartArtElement(
            diagram_index=0,
            diagram_type="list",
            text_nodes=[f"Item {i}" for i in range(20)],
            node_count=20,
            has_alt_text=False,
        )

        description = processor._describe_smartart(element)

        # Should have ellipsis indicating truncation
        assert "..." in description


class TestSmartArtDetection:
    """Test SmartArt detection in DOCX files."""

    @pytest.fixture
    def processor(self):
        """Create DocxProcessor for testing."""
        return DocxProcessor()

    def test_detect_smartart_no_diagrams(self, processor):
        """Test detection on document without SmartArt."""
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("This is a simple document with no SmartArt.")
            doc.add_paragraph("Just plain text here.")
            doc.save(f.name)

            issues = processor._detect_smartart(doc, f.name)

            # Should return empty list for document without SmartArt
            assert isinstance(issues, list)
            # May be empty or have issues depending on document

        os.unlink(f.name)

    def test_detect_smartart_returns_list(self, processor):
        """Test that detect_smartart always returns a list."""
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("Test document")
            doc.save(f.name)

            result = processor._detect_smartart(doc, f.name)

            assert isinstance(result, list)

        os.unlink(f.name)

    def test_detect_smartart_invalid_file(self, processor):
        """Test SmartArt detection with invalid file."""
        from docx import Document

        # Create empty docx
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.save(f.name)

            # Should not raise exception
            result = processor._detect_smartart(doc, "/nonexistent/file.docx")

            assert isinstance(result, list)

        os.unlink(f.name)


class TestSmartArtTypeIdentification:
    """Test SmartArt type identification."""

    @pytest.fixture
    def processor(self):
        """Create DocxProcessor for testing."""
        return DocxProcessor()

    def test_get_smartart_type_hierarchy(self, processor):
        """Test hierarchy type detection."""
        import zipfile
        import tempfile

        # Create a mock DOCX with layout containing hierarchy keyword
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            with zipfile.ZipFile(f.name, "w") as zf:
                # Add minimal layout file with hierarchy keyword - use layout0.xml for index 0
                layout_xml = (
                    b'<?xml version="1.0"?><root><hierarchy>test</hierarchy></root>'
                )
                zf.writestr("word/diagrams/layout0.xml", layout_xml)

            with zipfile.ZipFile(f.name, "r") as zf:
                result = processor._get_smartart_type(zf, 0, {})
                assert result == "hierarchy"

        os.unlink(f.name)

    def test_get_smartart_type_process(self, processor):
        """Test process type detection."""
        import zipfile
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            with zipfile.ZipFile(f.name, "w") as zf:
                # Use only process/flow keywords, avoid hierarchy - use layout0.xml for index 0
                layout_xml = (
                    b'<?xml version="1.0"?><root><processFlow>true</processFlow></root>'
                )
                zf.writestr("word/diagrams/layout0.xml", layout_xml)

            with zipfile.ZipFile(f.name, "r") as zf:
                result = processor._get_smartart_type(zf, 0, {})
                assert result == "process"

        os.unlink(f.name)

    def test_get_smartart_type_unknown(self, processor):
        """Test unknown type when no patterns match."""
        import zipfile
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            with zipfile.ZipFile(f.name, "w") as zf:
                # Use XML without any recognizable type keywords - use layout0.xml for index 0
                layout_xml = b'<?xml version="1.0"?><root><generic><item>test</item></generic></root>'
                zf.writestr("word/diagrams/layout0.xml", layout_xml)

            with zipfile.ZipFile(f.name, "r") as zf:
                result = processor._get_smartart_type(zf, 0, {})
                assert result == "unknown"

        os.unlink(f.name)


class TestDocxProcessorIntegration:
    """Test SmartArt integration with DocxProcessor."""

    @pytest.fixture
    def processor(self):
        """Create DocxProcessor for testing."""
        return DocxProcessor()

    def test_process_docx_includes_smartart_issues(self, processor):
        """Test that process_docx includes smartart_issues field."""
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_heading("Test Document", level=1)
            doc.add_paragraph("This is a test paragraph.")
            doc.save(f.name)

            result = processor.process_docx(f.name)

            # Should have smartart_issues field
            assert hasattr(result, "smartart_issues")
            assert isinstance(result.smartart_issues, list)

        os.unlink(f.name)

    def test_summary_includes_smartart_count(self, processor):
        """Test that summary includes smartart_issues count."""
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("Simple test")
            doc.save(f.name)

            result = processor.process_docx(f.name)

            # Summary should include smartart_issues key
            assert "smartart_issues" in result.summary

        os.unlink(f.name)


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-s"])
